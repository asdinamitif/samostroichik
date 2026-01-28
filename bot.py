import os
import re
import logging
import mimetypes
import hashlib
from pathlib import Path

URL_RE = re.compile(r"(https?://[^\s)\]\">]+)", re.IGNORECASE)


def extract_urls(*values: str) -> list[str]:
    urls: list[str] = []
    for v in values:
        if not v:
            continue
        s = str(v)
        for m in URL_RE.findall(s):
            u = m.strip().rstrip(".,;")
            if u not in urls:
                urls.append(u)
    return urls

def _token_for_url(url: str) -> str:
    return hashlib.sha1(url.encode("utf-8")).hexdigest()[:10]

def register_url(url: str, name_hint: str = "") -> str:
    t = _token_for_url(url)
    if t not in FILE_TOKEN_MAP:
        FILE_TOKEN_MAP[t] = {"url": url, "name": name_hint or ""}
    return t

def _guess_filename(url: str, fallback: str = "file") -> str:
    try:
        from urllib.parse import urlparse, unquote
        p = urlparse(url)
        seg = unquote(p.path.split("/")[-1])
        if seg and "." in seg:
            return seg
    except Exception:
        pass
    return fallback

def download_external_file(url: str) -> tuple[bytes, str, str]:
    """Скачивает файл по URL. Лимит по размеру: MAX_FILE_MB."""
    headers = {"User-Agent": "Mozilla/5.0 (compatible; SOTBot/1.0)"}
    r = requests.get(url, headers=headers, timeout=90, allow_redirects=True)
    r.raise_for_status()
    content = r.content
    size_mb = len(content) / (1024 * 1024)
    if size_mb > MAX_FILE_MB:
        raise RuntimeError(f"Файл слишком большой: {size_mb:.1f} MB (лимит {MAX_FILE_MB} MB)")
    fname = None
    cd = r.headers.get("content-disposition", "")
    m = re.search(r"filename\*=UTF-8''([^;]+)|filename=\"([^\"]+)\"|filename=([^;]+)", cd, flags=re.I)
    if m:
        fname = (m.group(1) or m.group(2) or m.group(3) or "").strip()
    fname = fname or _guess_filename(url, "document")
    mime = r.headers.get("content-type", "") or mimetypes.guess_type(fname)[0] or "application/octet-stream"
    return content, fname, mime

def analyze_file_bytes(data: bytes, filename: str, mime: str) -> str:
    """Лёгкий анализ содержимого: PDF/DOCX/TXT."""
    fn = filename.lower()
    if fn.endswith(".pdf") or "pdf" in (mime or ""):
        try:
            import PyPDF2
            from io import BytesIO
            reader = PyPDF2.PdfReader(BytesIO(data))
            text_parts = []
            for page in reader.pages[:10]:
                t = page.extract_text() or ""
                if t.strip():
                    text_parts.append(t)
            text = "\n".join(text_parts).strip()
            if not text:
                return "PDF загружен, но текст не извлекается (возможно, скан)."
            lines = [ln.strip() for ln in text.splitlines() if ln.strip()]
            return "Извлечённый текст (фрагмент):\n" + "\n".join(lines[:60])
        except Exception as e:
            return f"Не удалось разобрать PDF: {e}"

    if fn.endswith(".docx") or "wordprocessingml" in (mime or ""):
        try:
            from io import BytesIO
            import docx
            doc = docx.Document(BytesIO(data))
            paras = [p.text.strip() for p in doc.paragraphs if p.text and p.text.strip()]
            if not paras:
                return "DOCX загружен, но текста не найдено."
            return "Текст документа (фрагмент):\n" + "\n".join(paras[:80])
        except Exception as e:
            return f"Не удалось разобрать DOCX: {e}"

    if fn.endswith(".txt") or (mime or "").startswith("text/"):
        try:
            text = data.decode("utf-8", errors="ignore")
            lines = [ln.strip() for ln in text.splitlines() if ln.strip()]
            return "Текст (фрагмент):\n" + "\n".join(lines[:120])
        except Exception as e:
            return f"Не удалось разобрать текст: {e}"

    return "Файл скачан. Авто‑анализ доступен для PDF/DOCX/TXT. Могу прислать файл для просмотра."


import sqlite3
from datetime import datetime, timedelta, date
from io import BytesIO
from typing import Optional, Dict, Any, List, Tuple

import json
import requests
import pandas as pd
from openpyxl import load_workbook  # for reading xlsx with hyperlinks
from google.oauth2.service_account import Credentials
from googleapiclient.discovery import build
from dotenv import load_dotenv

import base64

# -------------------------------------------------
# YANDEX CLOUD (YandexGPT + SpeechKit) — ассистент
# -------------------------------------------------
def _get_env_any(*names: str) -> str:
    for n in names:
        v = os.getenv(n)
        if v is None:
            continue
        v = str(v).strip().strip('"').strip("'")
        if v:
            return v
    return ""

# Поддерживаем два нейминга переменных (старый YAGPT_* и новый YANDEX_*)
YANDEX_FOLDER_ID = _get_env_any("YANDEX_FOLDER_ID", "YAGPT_FOLDER_ID")
YANDEX_API_KEY = _get_env_any("YANDEX_API_KEY", "YAGPT_API_KEY")


ENABLE_ASSISTANT = os.getenv("ENABLE_ASSISTANT", "1").strip() == "1"
ENABLE_ASSISTANT_VOICE_REPLY = os.getenv("ENABLE_ASSISTANT_VOICE_REPLY", "1").strip() == "1"

logger = logging.getLogger(__name__)
logger.info("[YANDEX] folder_id_set=%s api_key_set=%s folder_id=%s",
            bool(YANDEX_FOLDER_ID), bool(YANDEX_API_KEY),
            (YANDEX_FOLDER_ID[:6] + "..." + YANDEX_FOLDER_ID[-4:]) if YANDEX_FOLDER_ID else "")

# Google Sheet для консультации (по умолчанию — то, что вы дали)
CONSULT_SHEET_URL = os.getenv(
    "CONSULT_SHEET_URL",
    "https://docs.google.com/spreadsheets/d/1W_9Cs-LaX6KR4cE9xN71CliE6Lm_TyQqk8t3kQa4FCc/edit?gid=967461758",
).strip()

def _extract_spreadsheet_id_from_url(url: str) -> str:
    if not url:
        return ""
    m = re.search(r"/d/([a-zA-Z0-9-_]+)", url)
    return m.group(1) if m else ""

def _extract_gid_from_url(url: str) -> str:
    if not url:
        return ""
    m = re.search(r"[?&]gid=(\d+)", url)
    return m.group(1) if m else ""

CONSULT_SHEET_ID = os.getenv("CONSULT_SHEET_ID", _extract_spreadsheet_id_from_url(CONSULT_SHEET_URL)).strip()
CONSULT_SHEET_GID = os.getenv("CONSULT_SHEET_GID", _extract_gid_from_url(CONSULT_SHEET_URL)).strip()

CONSULT_SHEET_CACHE_TTL_SEC = int(os.getenv("CONSULT_SHEET_CACHE_TTL_SEC", "600"))
CONSULT_SHEET_CACHE_PATH = os.getenv("CONSULT_SHEET_CACHE_PATH", "consult_sheet_cache.xlsx").strip()

# ==== FILE LINKS (download/analyze) ====
FILES_CACHE_DIR = Path(os.getenv('FILES_CACHE_DIR', '/data/files_cache')).resolve()
FILES_CACHE_DIR.mkdir(parents=True, exist_ok=True)
MAX_FILE_MB = int(os.getenv('MAX_FILE_MB', '25'))
FILE_TOKEN_MAP = {}  # token -> {'url':..., 'name':...}


def _now_ts() -> float:
    return datetime.utcnow().timestamp()

def download_google_sheet_xlsx(spreadsheet_id: str, gid: str, out_path: str) -> None:
    """
    Скачивает Google Sheet как XLSX через export. Для публичных/доступных по ссылке таблиц.
    """
    if not spreadsheet_id:
        raise ValueError("CONSULT_SHEET_ID пустой")
    url = f"https://docs.google.com/spreadsheets/d/{spreadsheet_id}/export?format=xlsx"
    if gid:
        url += f"&gid={gid}"
    r = requests.get(url, timeout=60)
    r.raise_for_status()
    Path(out_path).write_bytes(r.content)

def get_consult_df() -> pd.DataFrame:
    """
    Кэшируем XLSX локально, чтобы не дёргать Google на каждый вопрос.
    """
    cache = Path(CONSULT_SHEET_CACHE_PATH)
    meta = cache.with_suffix(".meta.json")
    use_cache = False
    if cache.exists() and meta.exists():
        try:
            m = json.loads(meta.read_text(encoding="utf-8"))
            ts = float(m.get("ts", 0))
            if _now_ts() - ts < CONSULT_SHEET_CACHE_TTL_SEC:
                use_cache = True
        except Exception:
            use_cache = False

    if not use_cache:
        download_google_sheet_xlsx(CONSULT_SHEET_ID, CONSULT_SHEET_GID, str(cache))
        meta.write_text(json.dumps({"ts": _now_ts()}, ensure_ascii=False), encoding="utf-8")

    # читаем первый лист по умолчанию
    return pd.read_excel(cache, sheet_name=0)

CASE_NO_RE = re.compile(r"\b\d{2}[-\s]?\d{2}[-\s]?\d{6}\b")

def normalize_case_no(s: str) -> Optional[str]:
    """
    Приводит номер дела к формату 00-00-000000.

    Поддерживает варианты:
    - '03-46-108600'
    - '03 46 108600'
    - '03 46 108 600'
    - '0 3 4 6 1 0 8 6 0 0' (STT часто так отдаёт)
    """
    if not s:
        return None
    s = str(s).strip()
    if not s:
        return None

    # 1) Ищем уже похожий на номер дела шаблон (с дефисами/пробелами)
    m = CASE_NO_RE.search(s.replace("—", "-").replace("–", "-"))
    if m:
        return f"{m.group(1)}-{m.group(2)}-{m.group(3)}"

    # 2) Фолбэк: вытаскиваем цифры и пробуем собрать 2-2-6
    digits = re.sub(r"\D+", "", s)
    if len(digits) == 10:
        return f"{digits[0:2]}-{digits[2:4]}-{digits[4:10]}"

    # Иногда в тексте могут быть лишние цифры (даты и т.п.) — попробуем найти 10-значный фрагмент
    if len(digits) > 10:
        for i in range(0, len(digits) - 9):
            cand = digits[i:i+10]
            # простая эвристика: первые 4 символа — это две группы по 2 цифры
            if re.match(r"^\d{10}$", cand):
                return f"{cand[0:2]}-{cand[2:4]}-{cand[4:10]}"

    return None
def safe_text(v) -> str:
    if v is None:
        return ""
    if isinstance(v, float) and pd.isna(v):
        return ""
    return str(v).strip()

def _yn_contains(val: str, target: str) -> bool:
    return target.lower() in safe_text(val).lower()

def _is_yes(val: str) -> bool:
    v = safe_text(val).lower()
    return v in ("да", "устранено", "true", "1", "yes", "y")

def _is_no(val: str) -> bool:
    v = safe_text(val).lower()
    return v in ("нет", "не устранено", "false", "0", "no", "n")

def find_case_rows(df: pd.DataFrame, case_no: str) -> pd.DataFrame:
    # пытаемся найти колонку "Номер дела" по заголовкам, иначе ищем по всем колонкам
    headers = [safe_text(c) for c in df.columns]
    idx_case = None
    for i, h in enumerate(headers):
        if "номер" in h.lower() and "дел" in h.lower():
            idx_case = i
            break
    if idx_case is not None:
        col = df.columns[idx_case]
        mask = df[col].astype(str).str.contains(case_no, na=False)
        out = df[mask].copy()
        if len(out) > 0:
            return out

    # fallback: по всем ячейкам
    mask_any = df.astype(str).apply(lambda row: row.str.contains(case_no, na=False)).any(axis=1)
    return df[mask_any].copy()

def pb_status_for_case(df: pd.DataFrame, case_no: str) -> Dict[str, Any]:
    """
    Возвращает статус устранения по ПБ.
    Приоритет: найти колонки по заголовкам (пожар/пб), иначе fallback на Q/R (как в вашей логике).
    """
    rows = find_case_rows(df, case_no)
    if rows.empty:
        return {"found": False, "case_no": case_no, "status": "NOT_FOUND"}

    headers = [safe_text(c) for c in df.columns]
    # поиск колонок по заголовкам
    pb_cols = []
    for i, h in enumerate(headers):
        hl = h.lower()
        if ("пожар" in hl) or (hl.startswith("пб")) or ("пб" in hl and "отмет" in hl):
            pb_cols.append(i)

    # fallback на Q/R (индексы 16/17) если колонок мало/не нашли
    if not pb_cols:
        # Q=17-я колонка (индекс 16), R=18-я (индекс 17)
        if len(headers) >= 18:
            pb_cols = [16, 17]

    if not pb_cols:
        return {"found": True, "case_no": case_no, "status": "NO_PB_COLUMNS", "details": {}}

    yes_seen = False
    no_seen = False
    details = {}
    for ci in pb_cols:
        colname = df.columns[ci]
        vals = [safe_text(v) for v in rows[colname].tolist()]
        details[safe_text(colname) or f"col_{ci}"] = vals
        for v in vals:
            if _is_no(v):
                no_seen = True
            if _is_yes(v):
                yes_seen = True

    if no_seen:
        st = "NOT_FIXED"
    elif yes_seen:
        st = "FIXED"
    else:
        st = "NO_DATA"

    return {"found": True, "case_no": case_no, "status": st, "details": details}

def retrieve_relevant_rows(df: pd.DataFrame, question: str, max_rows: int = 12) -> str:
    """
    Простая retrieval-стратегия: считаем совпадения токенов по строкам,
    возвращаем компактный контекст (заголовки + топ строк).
    """
    q = safe_text(question)
    tokens = [t for t in re.split(r"[\s,;:.!?()]+", q.lower()) if len(t) >= 3]
    if not tokens:
        tokens = [q.lower()] if q else []

    def score_row(row) -> int:
        s = " ".join([safe_text(x).lower() for x in row.values])
        return sum(1 for t in tokens if t in s)

    scored = []
    for i, row in df.iterrows():
        sc = score_row(row)
        if sc > 0:
            scored.append((sc, i))
    scored.sort(reverse=True)

    top_idx = [i for _, i in scored[:max_rows]]
    top = df.loc[top_idx] if top_idx else df.head(min(max_rows, len(df)))

    # ограничиваем длину
    cols = list(df.columns)
    header = " | ".join([safe_text(c) for c in cols])
    lines = [header]
    for _, r in top.iterrows():
        lines.append(" | ".join([safe_text(r[c]) for c in cols]))
    return "\n".join(lines[: max_rows + 1])

def yandex_chat_completion(system: str, user: str, temperature: float = 0.1, max_tokens: int = 800) -> str:
    """
    Yandex Cloud Foundation Models (YandexGPT) completion.

    Важно: у Yandex Cloud используется формат payload с modelUri и messages[text],
    а endpoint отличается от OpenAI.
    """
    if not YANDEX_API_KEY or not YANDEX_FOLDER_ID:
        raise RuntimeError("YANDEX_API_KEY / YANDEX_FOLDER_ID не заданы")
    url = "https://llm.api.cloud.yandex.net/foundationModels/v1/completion"
    headers = {
        "Authorization": f"Api-Key {YANDEX_API_KEY}",
        "Content-Type": "application/json",
        "x-folder-id": YANDEX_FOLDER_ID,
    }
    model_uri = os.getenv("YAGPT_MODEL") or f"gpt://{YANDEX_FOLDER_ID}/yandexgpt/latest"
    payload = {
        "modelUri": model_uri,
        "completionOptions": {
            "stream": False,
            "temperature": float(temperature),
            "maxTokens": int(max_tokens),
        },
        "messages": [
            {"role": "system", "text": system},
            {"role": "user", "text": user},
        ],
    }
    r = requests.post(url, headers=headers, json=payload, timeout=90)
    r.raise_for_status()
    data = r.json()
    # ожидаемый ответ: result.alternatives[0].message.text
    try:
        return safe_text(data["result"]["alternatives"][0]["message"]["text"]).strip()
    except Exception:
        # на случай вариаций ответа
        return safe_text(json.dumps(data, ensure_ascii=False))[:4000]

def yandex_speech_to_text(ogg_bytes: bytes) -> str:
    if not YANDEX_API_KEY or not YANDEX_FOLDER_ID:
        raise RuntimeError("YANDEX_API_KEY / YANDEX_FOLDER_ID не заданы")
    url = "https://stt.api.cloud.yandex.net/speech/v1/stt:recognize"
    headers = {
        "Authorization": f"Api-Key {YANDEX_API_KEY}",
        "Content-Type": "application/octet-stream",
        "x-folder-id": YANDEX_FOLDER_ID,
    }
    params = {
        "folderId": YANDEX_FOLDER_ID,
        "lang": "ru-RU",
        "format": "oggopus",
    }
    r = requests.post(url, headers=headers, params=params, data=ogg_bytes, timeout=90, allow_redirects=False)
    r.raise_for_status()
    data = r.json()
    return safe_text(data.get("result", "")).strip()

def yandex_text_to_speech(text_in: str) -> bytes:
    if not YANDEX_API_KEY or not YANDEX_FOLDER_ID:
        raise RuntimeError("YANDEX_API_KEY / YANDEX_FOLDER_ID не заданы")
    url = "https://tts.api.cloud.yandex.net/speech/v1/tts:synthesize"
    headers = {"Authorization": f"Api-Key {YANDEX_API_KEY}", "Content-Type": "application/octet-stream"}
    params = {
        "folderId": YANDEX_FOLDER_ID,
        "text": text_in,
        "lang": "ru-RU",
        "voice": "alena",
        "format": "oggopus",
        "speed": "1.0",
    }
    r = requests.post(url, headers=headers, data=params, timeout=90)
def _find_case_no_in_text(text: str) -> Optional[str]:
    # сначала — нормализация через общий парсер
    cn = normalize_case_no(text or "")
    if cn:
        return cn
    return None

# -------------------------------------------------
# Assistant: query parsing, normalization, export, snapshots
# -------------------------------------------------
SECTION_SYNONYMS = {
    "ПБ": ["пб", "пожар", "пожарка", "пожарной", "пожарная безопасность", "пожарная"],
    "АР": ["ар", "арх", "архитектура", "архитектур", "архитектурный"],
    "ММГН": ["ммгн", "мгн", "маломобильн", "доступ инвалид", "доступность", "инвалид"],
    "АГО": ["аго", "облик", "архитектурный облик", "арх облик"],
    "ЭОМ": ["эом", "электро", "электроснабж", "электрика"],
}

def _canon_section(text: str) -> Optional[str]:
    t = (text or "").lower()
    for canon, syns in SECTION_SYNONYMS.items():
        for s in syns:
            if s in t:
                return canon
    return None

def _find_case_no_in_text(text: str) -> Optional[str]:
    return extract_case_number(text)

def _col_by_keywords(df: pd.DataFrame, keywords: List[str]) -> Optional[str]:
    cols = list(df.columns)
    low = [str(c).strip().lower() for c in cols]
    for kw in keywords:
        kwl = kw.lower()
        for i, c in enumerate(low):
            if kwl in c:
                return cols[i]
    return None

def _status_columns(df: pd.DataFrame) -> Dict[str, str]:
    """
    Возвращает маппинг canonical_section -> column_name (если найдена)
    Поиск по заголовкам. Если не нашли – оставляем пусто.
    """
    m: Dict[str, str] = {}
    for canon, syns in SECTION_SYNONYMS.items():
        col = _col_by_keywords(df, [canon.lower()] + syns)
        if col:
            m[canon] = col
    return m

def _is_negative_status(v: str) -> bool:
    s = safe_text(v).strip().lower()
    if not s:
        return False
    negatives = ["нет", "не устран", "не устранено", "не выполн", "не выполнено", "0", "false", "❌"]
    return any(n in s for n in negatives)

def _is_positive_status(v: str) -> bool:
    s = safe_text(v).strip().lower()
    if not s:
        return False
    positives = ["да", "устран", "устранено", "выполн", "выполнено", "1", "true", "✅"]
    return any(p in s for p in positives) and not _is_negative_status(s)

def _universal_search(df: pd.DataFrame, query: str, limit: int = 30) -> pd.DataFrame:
    q = (query or "").strip()
    if not q:
        return df.head(0)
    ql = q.lower()
    mask = pd.Series([False] * len(df))
    for c in df.columns:
        try:
            mask = mask | df[c].astype(str).str.lower().str.contains(re.escape(ql), na=False)
        except Exception:
            continue
    return df[mask].head(limit)

def _df_case_card(df: pd.DataFrame, case_no: str) -> pd.DataFrame:
    if not case_no:
        return df.head(0)
    # ищем по всем колонкам, но приоритет на колонку "номер дела"
    col_case = _col_by_keywords(df, ["номер дела", "дело", "№ дела", "номер"])
    if col_case:
        d = df[df[col_case].astype(str).str.contains(re.escape(case_no), na=False)]
        if len(d) > 0:
            return d.head(1)
    # fallback – универсальный поиск по всем полям
    d2 = _universal_search(df, case_no, limit=1)
    return d2

def _df_not_fixed(df: pd.DataFrame, section: Optional[str] = None, limit: int = 50) -> pd.DataFrame:
    cols_map = _status_columns(df)
    if section and section in cols_map:
        col = cols_map[section]
        m = df[col].apply(_is_negative_status)
        return df[m].head(limit)
    # если раздел не распознан — ищем любые "нет" по всем колонкам статусов
    status_cols = list(cols_map.values())
    if not status_cols:
        return df.head(0)
    mask = pd.Series([False] * len(df))
    for c in status_cols:
        mask = mask | df[c].apply(_is_negative_status)
    return df[mask].head(limit)

def _key_fields(df: pd.DataFrame) -> Dict[str, Optional[str]]:
    return {
        "case": _col_by_keywords(df, ["номер дела", "дело", "№ дела", "номер"]),
        "address": _col_by_keywords(df, ["адрес", "местополож", "располож", "локац"]),
        "developer": _col_by_keywords(df, ["застрой", "организац", "ооо", "ао", "ип"]),
        "object": _col_by_keywords(df, ["объект", "наимен", "название"]),
    }

def _completeness_report(df: pd.DataFrame, limit: int = 80) -> pd.DataFrame:
    cols = _key_fields(df)
    status_cols = list(_status_columns(df).values())
    def is_empty(x) -> bool:
        s = safe_text(x).strip()
        return (not s) or s.lower() in ["nan", "none", "-", "—"]
    mask = pd.Series([False] * len(df))
    for k in ["case", "address", "developer", "object"]:
        c = cols.get(k)
        if c:
            mask = mask | df[c].apply(is_empty)
    for c in status_cols:
        mask = mask | df[c].apply(is_empty)
    return df[mask].head(limit)

def _hash_row(row: pd.Series) -> str:
    s = "|".join([safe_text(row.get(c)) for c in row.index])
    return hashlib.sha256(s.encode("utf-8", errors="ignore")).hexdigest()

def _snap_dir() -> Path:
    base = Path(DATA_DIR) if DATA_DIR else Path(".")
    d = base / "assistant_snapshots"
    d.mkdir(parents=True, exist_ok=True)
    return d

def save_snapshot(df: pd.DataFrame) -> Path:
    """
    Сохраняем снимок таблицы (хэши строк) на текущую дату.
    """
    cols = _key_fields(df)
    col_case = cols.get("case")
    snap = {}
    for _, row in df.iterrows():
        key = safe_text(row.get(col_case)) if col_case else ""
        key = extract_case_number(key) or safe_text(row.get(col_case)) or ""
        if not key:
            continue
        snap[key] = {"h": _hash_row(row)}
    p = _snap_dir() / f"{datetime.utcnow().date().isoformat()}.json"
    p.write_text(json.dumps({"ts": _now_ts(), "data": snap}, ensure_ascii=False), encoding="utf-8")
    return p

def _load_snapshot_near(days_ago: int = 7) -> Optional[dict]:
    target = (datetime.utcnow().date() - timedelta(days=days_ago))
    d = _snap_dir()
    if not d.exists():
        return None
    files = sorted(d.glob("*.json"))
    if not files:
        return None
    # выбираем файл с датой максимально близко к target (не позже target)
    best = None
    best_date = None
    for f in files:
        try:
            dt = datetime.strptime(f.stem, "%Y-%m-%d").date()
        except Exception:
            continue
        if dt <= target and (best_date is None or dt > best_date):
            best = f; best_date = dt
    if best is None:
        # fallback – самый ранний доступный
        best = files[0]
    try:
        return json.loads(best.read_text(encoding="utf-8"))
    except Exception:
        return None

def diff_week(df: pd.DataFrame) -> Dict[str, List[str]]:
    """
    Сравнение текущего состояния с снимком ~7 дней назад.
    Возвращает списки: new, removed, changed (case numbers).
    """
    prev = _load_snapshot_near(7)
    if not prev:
        # если нет снимков — сохраняем текущий и говорим, что сравнение будет доступно позже
        save_snapshot(df)
        return {"new": [], "removed": [], "changed": [], "note": ["Снимок для сравнения создан. Повторите запрос через неделю."]}
    prev_data = prev.get("data", {}) or {}
    cols = _key_fields(df)
    col_case = cols.get("case")
    cur = {}
    for _, row in df.iterrows():
        key = safe_text(row.get(col_case)) if col_case else ""
        key = extract_case_number(key) or safe_text(row.get(col_case)) or ""
        if not key:
            continue
        cur[key] = {"h": _hash_row(row)}
    new = sorted([k for k in cur.keys() if k not in prev_data])
    removed = sorted([k for k in prev_data.keys() if k not in cur])
    changed = sorted([k for k in cur.keys() if k in prev_data and cur[k]["h"] != prev_data[k].get("h")])
    # сохраняем текущий снимок в любом случае
    save_snapshot(df)
    return {"new": new, "removed": removed, "changed": changed, "note": []}

def df_to_file_bytes(df: pd.DataFrame, fmt: str = "xlsx") -> Tuple[BytesIO, str]:
    fmt = (fmt or "xlsx").lower()
    bio = BytesIO()
    if fmt == "csv":
        df.to_csv(bio, index=False, encoding="utf-8-sig")
        name = "export.csv"
    else:
        with pd.ExcelWriter(bio, engine="openpyxl") as w:
            df.to_excel(w, index=False, sheet_name="export")
        name = "export.xlsx"
    bio.seek(0)
    bio.name = name
    return bio, name

async def assistant_answer(chat, context, question_text: str, recognized_from_voice: bool = False):
    """
    Главная точка ассистента: детерминированные ответы (например ПБ) + общий Q&A по таблице.
    """
    q = safe_text(question_text)
    if recognized_from_voice:
        await chat.send_message(f"🎙 Распознано: {q}")

    if not q:
        await chat.send_message("Не получилось распознать запрос. Повторите, пожалуйста.")
        return

    # быстрый выход
    if q.lower() in ("выход", "стоп", "назад"):
        context.user_data["assistant_mode"] = False
        await chat.send_message("Режим ассистента выключен.", reply_markup=main_menu())
        return

        # 1) детерминированные операции по таблице (поиск/карточка/статусы/выгрузка/контроль полноты/изменения)
        ql = q.lower()

        export_requested = any(w in ql for w in ["выгруз", "экспорт", "csv", "excel", "xlsx", "файл", "пришли файлом", "сформируй"])
        export_fmt = "csv" if "csv" in ql else "xlsx"

        completeness_requested = any(w in ql for w in ["проверка полноты", "полнота", "пустые поля", "не заполн", "пусто в", "пустые"])
        diff_requested = ("что измен" in ql or "изменени" in ql) and ("недел" in ql or "7 " in ql or "сем" in ql)

        section = _canon_section(q)
        case_no = _find_case_no_in_text(q)  # 00-00-000000

        try:
            df = get_consult_df()

            # 1.1) Сравнение за неделю (по снимкам)
            if diff_requested:
                d = diff_week(df)
                note = d.get("note", [])
                lines = []
                if note:
                    lines.extend(note)
                else:
                    lines.append("Изменения относительно снимка ~7 дней назад:")
                    lines.append(f"• Новые дела: {len(d['new'])}")
                    lines.append(f"• Удалены/пропали: {len(d['removed'])}")
                    lines.append(f"• Изменены строки: {len(d['changed'])}")
                    # показать первые 20 номеров
                    if d["new"]:
                        lines.append("Новые (первые 20): " + ", ".join(d["new"][:20]))
                    if d["changed"]:
                        lines.append("Изменены (первые 20): " + ", ".join(d["changed"][:20]))
                    if d["removed"]:
                        lines.append("Удалены (первые 20): " + ", ".join(d["removed"][:20]))
                out = "\n".join(lines)
                await chat.send_message(out)
                if ENABLE_ASSISTANT_VOICE_REPLY:
                    try:
                        audio = yandex_text_to_speech(out[:800])
                        bio = BytesIO(audio); bio.name = "answer.ogg"
                        await chat.send_voice(voice=bio)
                    except Exception:
                        pass
                return

            # 1.2) Проверка полноты (пустые ключевые поля/статусы)
            if completeness_requested:
                bad = _completeness_report(df, limit=80)
                if len(bad) == 0:
                    out = "✅ Пустых ключевых полей не найдено (по эвристике: адрес/застройщик/объект/статусы)."
                    await chat.send_message(out)
                else:
                    out = (
                        f"⚠️ Найдены строки с пустыми ключевыми полями/статусами: {len(bad)} (показываю до 80)."
                    )
                    await chat.send_message(out)
                    # Если просили выгрузку — сразу отправим файлом
                    if export_requested:
                        bio, fname = df_to_file_bytes(bad, fmt=export_fmt)
                        await chat.send_document(document=bio, filename=fname, caption="Выгрузка: пустые/не заполненные поля")
                    else:
                        # краткий список по номеру дела (если есть)
                        cols = _key_fields(df)
                        col_case = cols.get("case")
                        sample = []
                        if col_case and col_case in bad.columns:
                            for v in bad[col_case].astype(str).head(30).tolist():
                                cn = extract_case_number(v) or v
                                sample.append(cn)
                            await chat.send_message("Примеры (до 30): " + ", ".join(sample))
                return

            # 1.3) Если есть номер дела — карточка/проверка статуса раздела
            if case_no:
                row = _df_case_card(df, case_no)
                if len(row) == 0:
                    await chat.send_message(f"Дело {case_no}: не найдено в таблице.")
                    return

                # проверка конкретного раздела (ПБ/АР/ММГН/АГО/ЭОМ и т.п.)
                if section:
                    cols_map = _status_columns(df)
                    col = cols_map.get(section)
                    if not col:
                        await chat.send_message(f"По делу {case_no}: не нашёл колонку для раздела «{section}» в таблице.")
                        return
                    val = row.iloc[0].get(col)
                    if _is_positive_status(val):
                        out = f"✅ По делу {case_no}: по «{section}» устранено."
                    elif _is_negative_status(val):
                        out = f"❌ По делу {case_no}: по «{section}» НЕ устранено (есть «нет»/аналог)."
                    else:
                        out = f"ℹ️ По делу {case_no}: по «{section}» нет отметки/пусто."
                    await chat.send_message(out)
                    if export_requested:
                        bio, fname = df_to_file_bytes(row, fmt=export_fmt)
                        await chat.send_document(document=bio, filename=fname, caption=f"Выгрузка: дело {case_no}")
                    if ENABLE_ASSISTANT_VOICE_REPLY:
                        try:
                            audio = yandex_text_to_speech(out[:800])
                            bio = BytesIO(audio); bio.name = "answer.ogg"
                            await chat.send_voice(voice=bio)
                        except Exception:
                            pass
                    return

                # если раздел не указан — покажем карточку (коротко) + при желании выгрузку
                cols = _key_fields(df)
                cols_map = _status_columns(df)
                fields = []
                for k, label in [("case","Номер дела"),("address","Адрес"),("developer","Застройщик"),("object","Объект")]:
                    c = cols.get(k)
                    if c and c in row.columns:
                        fields.append(f"• {label}: {safe_text(row.iloc[0].get(c))}")
                # статусы
                for sec, col in cols_map.items():
                    if col in row.columns:
                        fields.append(f"• {sec}: {safe_text(row.iloc[0].get(col))}")
                out = "Карточка по делу:\n" + "\n".join(fields) if fields else f"Нашёл строку по делу {case_no}."
                await chat.send_message(out)
                if export_requested:
                    bio, fname = df_to_file_bytes(row, fmt=export_fmt)
                    await chat.send_document(document=bio, filename=fname, caption=f"Выгрузка: дело {case_no}")
                return

            # 1.4) Списки "не устранено" по разделу (или в целом по статусам)
            if any(w in ql for w in ["не устран", "неустран", "есть нет", "где нет", "статус нет"]):
                bad = _df_not_fixed(df, section=section, limit=60)
                if len(bad) == 0:
                    await chat.send_message("По запросу «не устранено» ничего не найдено.")
                    return
                title = f"❌ Не устранено по разделу «{section}»" if section else "❌ Не устранено (по доступным статус-колонкам)"
                await chat.send_message(f"{title}\nНайдено: {len(bad)} (показываю до 60).")
                if export_requested:
                    bio, fname = df_to_file_bytes(bad, fmt=export_fmt)
                    await chat.send_document(document=bio, filename=fname, caption="Выгрузка: не устранено")
                else:
                    cols = _key_fields(df)
                    col_case = cols.get("case")
                    col_addr = cols.get("address")
                    lines = []
                    for _, r in bad.head(30).iterrows():
                        cn = extract_case_number(safe_text(r.get(col_case))) if col_case else ""
                        addr = safe_text(r.get(col_addr)) if col_addr else ""
                        if cn or addr:
                            lines.append(f"• {cn or '—'} — {addr[:120]}")
                    if lines:
                        await chat.send_message("\n".join(lines))
                return

            # 1.5) Универсальный поиск по любому полю + выгрузка
            # (например: "найди по застройщику ...", "найди адрес ...", и т.п.)
            if ql.startswith("найди") or ql.startswith("покажи") or export_requested:
                # извлечём "хвост" после "найди/покажи"
                q2 = q
                for pfx in ["найди", "покажи", "поиск", "сформируй", "выгрузка", "экспорт"]:
                    if ql.startswith(pfx):
                        q2 = q[len(pfx):].strip(" :,-")
                        break
                res = _universal_search(df, q2 or q, limit=200 if export_requested else 30)
                if len(res) == 0:
                    await chat.send_message("Ничего не найдено по запросу.")
                    return
                await chat.send_message(f"Найдено строк: {len(res)} (показываю до {len(res)}).")
                if export_requested:
                    bio, fname = df_to_file_bytes(res, fmt=export_fmt)
                    await chat.send_document(document=bio, filename=fname, caption="Выгрузка по запросу")
                else:
                    cols = _key_fields(df)
                    col_case = cols.get("case")
                    col_addr = cols.get("address")
                    lines = []
                    for _, r in res.head(25).iterrows():
                        cn = extract_case_number(safe_text(r.get(col_case))) if col_case else ""
                        addr = safe_text(r.get(col_addr)) if col_addr else ""
                        if cn or addr:
                            lines.append(f"• {cn or '—'} — {addr[:120]}")
                    if lines:
                        await chat.send_message("\n".join(lines))
                return

            # если ничего не сработало — продолжаем в общий Q&A
        except Exception as e:
            await chat.send_message(f"Ошибка при обработке по таблице: {e}")
            return

# 2) общий Q&A: retrieval по таблице + YandexGPT
    try:
        df = get_consult_df()
        context_rows = retrieve_relevant_rows(df, q, max_rows=12)
        system = (
            "Ты — служебный ассистент Главгосстройнадзора. "
            "Отвечай строго по предоставленному фрагменту таблицы. "
            "Если данных недостаточно — скажи, что в таблице нет сведений, и уточни, что нужно."
        )
        user_prompt = (
            f"Вопрос пользователя:\n{q}\n\n"
            f"Фрагмент таблицы (заголовки и строки):\n{context_rows}\n\n"
            "Дай краткий, точный ответ. Если можно — укажи номер дела и ключевые поля."
        )
        answer = yandex_chat_completion(system, user_prompt, temperature=0.1, max_tokens=700)
        await chat.send_message(answer)

        if ENABLE_ASSISTANT_VOICE_REPLY:
            try:
                audio = yandex_text_to_speech(answer[:800])
                bio = BytesIO(audio); bio.name = "answer.ogg"
                await chat.send_voice(voice=bio)
            except Exception:
                pass

    except Exception as e:
        await chat.send_message(
            "Ассистент недоступен (проверьте YANDEX_API_KEY/YANDEX_FOLDER_ID и доступ к таблице). "
            f"Детали: {e}"
        )

from telegram import (
    Update,
    ReplyKeyboardMarkup,
    InlineKeyboardMarkup,
    InlineKeyboardButton,
    InputFile,
)
from telegram.ext import (
    Application,
    CommandHandler,
    MessageHandler,
    CallbackQueryHandler,
    ContextTypes,
    filters,
)

from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.worksheet.table import Table, TableStyleInfo

AnyType = Any

# ----------------- ЛОГИ -----------------
logging.basicConfig(
    format="%(asctime)s [%(levelname)s] %(name)s: %(message)s",
    level=logging.INFO,
)
log = logging.getLogger("sot_bot")

# ----------------- НАСТРОЙКИ И .ENV -----------------
load_dotenv()

BOT_TOKEN = (os.getenv("BOT_TOKEN") or "").strip()
DB_PATH = os.getenv("DB_PATH", "sot_bot.db")

TIMEZONE_OFFSET = int(os.getenv("TIMEZONE_OFFSET", "3"))
ANALYTICS_PASSWORD = "051995"

# ----------------- ДАННЫЕ / PERSISTENCE (Railway Volume) -----------------
# В Railway обязательно подключить Volume и примонтировать в /data,
# чтобы результаты и пароль сохранялись между перезапусками.
DATA_DIR = (os.getenv("DATA_DIR") or "/data").strip() or "/data"
os.makedirs(DATA_DIR, exist_ok=True)

# ----------------- ШАБЛОН DOCX (ТЗ для ЦНИЛ) -----------------
# Можно задать переменную окружения CNIL_T3_TEMPLATE.
# Иначе берём стандартный путь в репозитории (/app) или в Volume (/data).
CNIL_T3_TEMPLATE = (os.getenv("CNIL_T3_TEMPLATE") or os.getenv("TEST_T3_TEMPLATE") or "/app/TEST_T3.docx").strip()

def _resolve_cnil_t3_template(path: str) -> str:
    candidates = [
        path,
        "/app/TEST_T3.docx",
        "/app/TEST_T3_TEMPLATE.docx",
        os.path.join(DATA_DIR, "TEST_T3.docx"),
        os.path.join(DATA_DIR, "TEST_T3_TEMPLATE.docx"),
    ]
    for c in candidates:
        try:
            if c and os.path.exists(c):
                return c
        except Exception:
            pass
    return path

CNIL_T3_TEMPLATE = _resolve_cnil_t3_template(CNIL_T3_TEMPLATE)


# =========================
# 🧪 ТЗ для ЦНИЛ: пароль на скачивание
# =========================

# Резервный пароль (НЕ меняется никогда)
CNIL_MASTER_DOWNLOAD_PASSWORD = "051995"

def cnil_password_file() -> str:
    return os.path.join(DATA_DIR, "cnil_download_password.json")

def cnil_load_download_password() -> str:
    """Текущий пароль на скачивание. По умолчанию: 1234."""
    path = cnil_password_file()
    try:
        if os.path.exists(path):
            with open(path, "r", encoding="utf-8") as f:
                obj = json.load(f)
            pw = str(obj.get("password", "")).strip()
            return pw or "1234"
    except Exception:
        pass
    return "1234"

def cnil_save_download_password(new_password: str) -> None:
    path = cnil_password_file()
    obj = {"password": str(new_password).strip()}
    tmp = path + ".tmp"
    with open(tmp, "w", encoding="utf-8") as f:
        json.dump(obj, f, ensure_ascii=False)
    os.replace(tmp, path)


def _extract_spreadsheet_id_from_url(url: str) -> str:
    try:
        if "/d/" in url:
            return url.split("/d/")[1].split("/")[0]
    except Exception:
        pass
    return ""


SCHEDULE_URL_ENV = (os.getenv("SCHEDULE_URL") or "").strip()

_default_sheet_id = _extract_spreadsheet_id_from_url(SCHEDULE_URL_ENV)
if not _default_sheet_id:
    _default_sheet_id = (os.getenv("GSHEETS_SPREADSHEET_ID") or "").strip()
if not _default_sheet_id:
    _default_sheet_id = "1W_9Cs-LaX6KR4cE9xN71CliE6Lm_TyQqk8t3kQa4FCc"

GSHEETS_SPREADSHEET_ID = _default_sheet_id

if SCHEDULE_URL_ENV:
    GOOGLE_SHEET_URL_DEFAULT = SCHEDULE_URL_ENV
else:
    GOOGLE_SHEET_URL_DEFAULT = (
        f"https://docs.google.com/spreadsheets/d/{GSHEETS_SPREADSHEET_ID}/edit?usp=sharing"
    )

GSHEETS_SERVICE_ACCOUNT_JSON = (os.getenv("GSHEETS_SERVICE_ACCOUNT_JSON") or "").strip()
SHEETS_SERVICE = None

DEFAULT_APPROVERS = [
    "@asdinamitif",
    "@FrolovAlNGSN",
    "@cappit_G59",
    "@sergeybektiashkin",
    "@scri4",
    "@Kirill_Victorovi4",
]

RESPONSIBLE_USERNAMES: Dict[str, List[str]] = {
    "бектяшкин": ["sergeybektiashkin"],
    "смирнов": ["scri4"],
}

INSPECTOR_SHEET_NAME = "ПБ, АР,ММГН, АГО (2025)"
HARD_CODED_ADMINS = {398960707}

SCHEDULE_NOTIFY_CHAT_ID_ENV = (os.getenv("SCHEDULE_NOTIFY_CHAT_ID") or "").strip()
SCHEDULE_NOTIFY_CHAT_ID = (
    int(SCHEDULE_NOTIFY_CHAT_ID_ENV) if SCHEDULE_NOTIFY_CHAT_ID_ENV else None
)

# ВТОРАЯ ТАБЛИЦА — итоговые проверки
FINAL_CHECKS_SPREADSHEET_ID = (
    os.getenv(
        "FINAL_CHECKS_SPREADSHEET_ID",
        "1dUO3neTKzKI3D8P6fs_LJLmWlL7jw-FhohtJkjz4KuE",
    ).strip()
)


FINAL_CHECKS_LOCAL_PATH = os.getenv(
    "FINAL_CHECKS_LOCAL_PATH",
    "final_checks.xlsx",
).strip()



def is_admin(uid: int) -> bool:
    return uid in HARD_CODED_ADMINS


def local_now() -> datetime:
    return datetime.utcnow() + timedelta(hours=TIMEZONE_OFFSET)


def get_current_remarks_sheet_name() -> str:
    year = local_now().year
    return f"ПБ, АР,ММГН, АГО ({year})"


# -------------------------------------------------
# Google Sheets helpers
# -------------------------------------------------
def get_sheets_service():
    global SHEETS_SERVICE

    if SHEETS_SERVICE is not None:
        return SHEETS_SERVICE

    if not GSHEETS_SERVICE_ACCOUNT_JSON:
        log.error(
            "GSHEETS_SERVICE_ACCOUNT_JSON не задан – Google Sheets API недоступен."
        )
        return None

    try:
        info = json.loads(GSHEETS_SERVICE_ACCOUNT_JSON)
        creds = Credentials.from_service_account_info(
            info,
            scopes=["https://www.googleapis.com/auth/spreadsheets"],
        )
        service = build("sheets", "v4", credentials=creds)
        SHEETS_SERVICE = service
        return service
    except Exception as e:
        log.error("Ошибка создания клиента Google Sheets: %s", e)
        return None


def build_export_url(spreadsheet_id: str) -> str:
    return f"https://docs.google.com/spreadsheets/d/{spreadsheet_id}/export?format=xlsx"


def detect_header_row(values: List[List[str]]) -> int:
    for i, row in enumerate(values[:30]):
        row_lower = [str(c).lower() for c in row]
        if any("дата выезда" in c for c in row_lower):
            return i
    return 0


def read_sheet_to_dataframe(
    sheet_id: str, sheet_name: str, header_row_index: Optional[int] = None
) -> Optional[pd.DataFrame]:
    service = get_sheets_service()
    if service is None:
        log.error("Google Sheets сервис недоступен – невозможно прочитать лист.")
        return None

    try:
        result = (
            service.spreadsheets()
            .values()
            .get(spreadsheetId=sheet_id, range=f"'{sheet_name}'!A1:ZZZ1000")
            .execute()
        )
        values = result.get("values", [])

        if not values:
            log.warning("Лист '%s' пуст.", sheet_name)
            return pd.DataFrame()

        if header_row_index is None:
            header_row_index = detect_header_row(values)

        headers = values[header_row_index]
        data_rows = values[header_row_index + 1 :]

        df = pd.DataFrame(data_rows, columns=headers)
        df = df.dropna(how="all").reset_index(drop=True)
        return df
    except Exception as e:
        log.error("Ошибка чтения листа '%s' из Google Sheets: %s", sheet_name, e)
        return None


# -------------------------------------------------
# Работа со столбцами Excel
# -------------------------------------------------
def excel_col_to_index(col: str) -> int:
    col = col.upper().strip()
    idx = 0
    for ch in col:
        if "A" <= ch <= "Z":
            idx = idx * 26 + (ord(ch) - ord("A") + 1)
    return idx - 1


def get_col_by_letter(df: pd.DataFrame, letters: str) -> Optional[str]:
    idx = excel_col_to_index(letters)
    if 0 <= idx < len(df.columns):
        return df.columns[idx]
    return None


def get_col_index_by_header(
    df: pd.DataFrame, search_substr: str, fallback_letter: str
) -> Optional[int]:
    search_substr = search_substr.lower()
    for i, col in enumerate(df.columns):
        if search_substr in str(col).lower():
            return i
    idx = excel_col_to_index(fallback_letter)
    if 0 <= idx < len(df.columns):
        return idx
    return None


def normalize_onzs_value(val) -> Optional[str]:
    if val is None:
        return None
    s = str(val).strip()
    if not s:
        return None
    try:
        n = int(float(s.replace(",", ".")))
        return str(n)
    except Exception:
        pass
    return s


def normalize_case_number(val) -> str:
    """
    Нормализация номера дела:

    - приводим все нестандартные тире к обычному '-';
    - убираем пробелы;
    - выбрасываем любые символы, кроме цифр и '-'.

    Примеры:
    'Дело № 03–46–108600 (ПП)' -> '03-46-108600'
    ' 01-29-099900 ' -> '01-29-099900'
    """
    if val is None:
        return ""
    s = str(val).strip()
    if not s:
        return ""

    # все «косые» тире в нормальное
    hyphens = ["\u2010", "\u2011", "\u2012", "\u2013", "\u2014", "\u2212"]
    for h in hyphens:
        s = s.replace(h, "-")

    # убираем пробелы
    s = s.replace(" ", "")

    # оставляем только цифры и '-'
    cleaned_chars = []
    for ch in s:
        if ch.isdigit() or ch == "-":
            cleaned_chars.append(ch)

    return "".join(cleaned_chars)


def get_case_col_index(df: pd.DataFrame) -> Optional[int]:
    idx_i = excel_col_to_index("I")
    if 0 <= idx_i < len(df.columns):
        return idx_i
    return get_col_index_by_header(df, "номер дела", "I")


# -------------------------------------------------
# БАЗА ДАННЫХ
# -------------------------------------------------
def get_db() -> sqlite3.Connection:
    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row
    return conn


def init_db() -> None:
    conn = get_db()
    c = conn.cursor()

    c.execute(
        """CREATE TABLE IF NOT EXISTS schedule_settings (
               key TEXT PRIMARY KEY,
               value TEXT
           )"""
    )

    c.execute(
        """CREATE TABLE IF NOT EXISTS approvers (
               id INTEGER PRIMARY KEY AUTOINCREMENT,
               label TEXT UNIQUE
           )"""
    )

    c.execute(
        """CREATE TABLE IF NOT EXISTS schedule_files (
               version INTEGER PRIMARY KEY,
               name TEXT,
               uploaded_at TEXT
           )"""
    )

    c.execute(
        """CREATE TABLE IF NOT EXISTS schedule_approvals (
               id INTEGER PRIMARY KEY AUTOINCREMENT,
               version INTEGER,
               approver TEXT,
               status TEXT,
               comment TEXT,
               decided_at TEXT,
               requested_at TEXT
           )"""
    )

    c.execute(
        """CREATE TABLE IF NOT EXISTS inspector_visits (
               id INTEGER PRIMARY KEY AUTOINCREMENT,
               date TEXT,
               area TEXT,
               floors TEXT,
               onzs TEXT,
               developer TEXT,
               object TEXT,
               address TEXT,
               case_no TEXT,
               check_type TEXT,
               created_at TEXT
           )"""
    )

    c.execute("SELECT COUNT(*) AS c FROM approvers")
    if c.fetchone()["c"] == 0:
        c.executemany(
            "INSERT OR IGNORE INTO approvers (label) VALUES (?)",
            [(lbl,) for lbl in DEFAULT_APPROVERS],
        )

    c.execute("SELECT value FROM schedule_settings WHERE key='schedule_version'")
    if not c.fetchone():
        c.execute(
            "INSERT INTO schedule_settings (key, value) VALUES ('schedule_version', '1')"
        )

    c.execute("SELECT value FROM schedule_settings WHERE key='last_notified_version'")
    if not c.fetchone():
        c.execute(
            "INSERT INTO schedule_settings (key, value) VALUES ('last_notified_version', '0')"
        )

    if SCHEDULE_NOTIFY_CHAT_ID_ENV:
        c.execute(
            "INSERT OR REPLACE INTO schedule_settings (key, value) VALUES (?, ?)",
            ("schedule_notify_chat_id", SCHEDULE_NOTIFY_CHAT_ID_ENV),
        )

    conn.commit()
    conn.close()


def get_schedule_state() -> dict:
    conn = get_db()
    c = conn.cursor()
    c.execute("SELECT key, value FROM schedule_settings")
    rows = c.fetchall()
    conn.close()
    return {r["key"]: r["value"] for r in rows}


def get_schedule_version(settings: dict) -> int:
    try:
        return int(settings.get("schedule_version") or "1")
    except Exception:
        return 1


def get_current_approvers(settings: dict) -> List[str]:
    val = settings.get("current_approvers")
    if val:
        arr = [v.strip() for v in val.split(",") if v.strip()]
        if arr:
            return arr
    return []


def set_current_approvers_for_version(approvers: List[str], version: int) -> None:
    conn = get_db()
    c = conn.cursor()

    c.execute(
        "INSERT OR REPLACE INTO schedule_settings (key, value) VALUES ('current_approvers', ?)",
        (",".join(approvers),),
    )

    c.execute("DELETE FROM schedule_approvals WHERE version = ?", (version,))

    now = local_now().isoformat()
    for appr in approvers:
        c.execute(
            """INSERT INTO schedule_approvals
               (version, approver, status, comment, decided_at, requested_at)
               VALUES (?, ?, 'pending', NULL, NULL, ?)""",
            (version, appr, now),
        )

    conn.commit()
    conn.close()


def get_schedule_approvals(version: int) -> List[sqlite3.Row]:
    conn = get_db()
    c = conn.cursor()
    c.execute(
        "SELECT * FROM schedule_approvals WHERE version = ? ORDER BY approver",
        (version,),
    )
    rows = c.fetchall()
    conn.close()
    return rows


def update_schedule_approval_status(
    version: int, approver: str, status: str, comment: Optional[str] = None
):
    conn = get_db()
    c = conn.cursor()
    now = local_now().isoformat()

    c.execute(
        """UPDATE schedule_approvals
           SET status=?, comment=?, decided_at=?
           WHERE version=? AND approver=?""",
        (status, comment, now, version, approver),
    )
    conn.commit()
    conn.close()


# -------------------------------------------------
# Инспектор: БД
# -------------------------------------------------
def save_inspector_to_db(form: Dict[str, Any]) -> bool:
    try:
        conn = get_db()
        c = conn.cursor()
        date_obj = form.get("date")
        date_str = date_obj.strftime("%Y-%m-%d") if date_obj else None
        c.execute(
            """INSERT INTO inspector_visits
               (date, area, floors, onzs, developer, object, address,
                case_no, check_type, created_at)
               VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?)""",
            (
                date_str,
                form.get("area", ""),
                form.get("floors", ""),
                form.get("onzs", ""),
                form.get("developer", ""),
                form.get("object", ""),
                form.get("address", ""),
                form.get("case", ""),
                form.get("check_type", ""),
                local_now().isoformat(),
            ),
        )
        conn.commit()
        conn.close()
        return True
    except Exception as e:
        log.error("Ошибка сохранения инспектора в локную БД: %s", e)
        return False


def fetch_inspector_visits(limit: int = 50) -> List[sqlite3.Row]:
    conn = get_db()
    c = conn.cursor()
    c.execute(
        """SELECT * FROM inspector_visits
           ORDER BY date DESC, id DESC
           LIMIT ?""",
        (limit,),
    )
    rows = c.fetchall()
    conn.close()
    return rows


def clear_inspector_visits() -> None:
    conn = get_db()
    c = conn.cursor()
    c.execute("DELETE FROM inspector_visits")
    conn.commit()
    conn.close()


# -------------------------------------------------
# Клавиатуры
# -------------------------------------------------
def main_menu() -> ReplyKeyboardMarkup:
    return ReplyKeyboardMarkup(
        [
            ["📅 График", "📝 Замечания"],
            ["Инспектор", "📈 Аналитика"],
            ["Итоговые проверки"],
            ["🚨 Красные лампочки"],
            ["🧪 ТЗ для ЦНИЛ"],
            ["🗣 Ассистент"],
        ],
        resize_keyboard=True,
    )


def build_schedule_inline(
    is_admin_flag: bool, settings: dict, user_tag: Optional[str] = None
) -> InlineKeyboardMarkup:
    buttons = [
        [
            InlineKeyboardButton("🔄 Обновить", callback_data="schedule_refresh"),
            InlineKeyboardButton("📥 Скачать", callback_data="schedule_download"),
        ],
        [InlineKeyboardButton("📤 Загрузить", callback_data="schedule_upload")],
    ]
    if is_admin_flag:
        buttons.append(
            [InlineKeyboardButton("👥 Согласующие", callback_data="schedule_approvers")]
        )

    approvers = get_current_approvers(settings)
    if user_tag and user_tag in approvers:
        buttons.append(
            [
                InlineKeyboardButton(
                    f"✅ Согласовать ({user_tag})",
                    callback_data=f"schedule_approve:{user_tag}",
                ),
                InlineKeyboardButton(
                    f"✏️ На доработку ({user_tag})",
                    callback_data=f"schedule_rework:{user_tag}",
                ),
            ]
        )

    return InlineKeyboardMarkup(buttons)


def remarks_menu_inline() -> InlineKeyboardMarkup:
    return InlineKeyboardMarkup(
        [
            [
                InlineKeyboardButton(
                    "🔎 Поиск по номеру дела", callback_data="remarks_search_case"
                )
            ],
            [InlineKeyboardButton("🏗 ОНзС", callback_data="remarks_onzs")],
            [InlineKeyboardButton("📥 Открыть файл", callback_data="remarks_download")],
        ]
    )


def inspector_menu_inline() -> InlineKeyboardMarkup:
    return InlineKeyboardMarkup(
        [
            [InlineKeyboardButton("➕ Добавить выезд", callback_data="inspector_add")],
            [
                InlineKeyboardButton("📋 Список выездов", callback_data="inspector_list"),
                InlineKeyboardButton(
                    "📥 Скачать Excel", callback_data="inspector_download"
                ),
            ],
            [
                InlineKeyboardButton(
                    "🔄 Обновить", callback_data="inspector_reset"
                )
            ],
        ]
    )


def final_checks_menu_inline() -> InlineKeyboardMarkup:
    return InlineKeyboardMarkup(
        [
            [
                InlineKeyboardButton("📅 За неделю", callback_data="final_week"),
                InlineKeyboardButton("📆 За месяц", callback_data="final_month"),
            ],
            [
                InlineKeyboardButton(
                    "📊 Выбрать период", callback_data="final_period"
                )
            ],
            [
                InlineKeyboardButton(
                    "🔎 По номеру дела", callback_data="final_search_case"
                )
            ],
        ]
    )


# -------------------------------------------------
# График
# -------------------------------------------------
def get_schedule_df() -> Optional[pd.DataFrame]:
    SHEET = "График"
    url = build_export_url(GSHEETS_SPREADSHEET_ID)

    try:
        resp = requests.get(url, timeout=30)
        resp.raise_for_status()
    except Exception as e:
        log.error("Ошибка скачивания Excel (график): %s", e)
        return None

    try:
        xls = pd.ExcelFile(BytesIO(resp.content))
        if SHEET not in xls.sheet_names:
            log.error("В файле нет листа '%s'", SHEET)
            return None
        df = pd.read_excel(xls, sheet_name=SHEET)
        df = df.dropna(how="all").reset_index(drop=True)
        return df
    except Exception as e:
        log.error("Ошибка чтения листа графика: %s", e)
        return None


HEADER_FILL = PatternFill(start_color="305496", end_color="305496", fill_type="solid")
HEADER_FONT = Font(color="FFFFFF", bold=True)
BORDER = Border(
    left=Side(style="thin"),
    right=Side(style="thin"),
    top=Side(style="thin"),
    bottom=Side(style="thin"),
)


async def send_schedule_xlsx(
    chat_id: int, dataframe: pd.DataFrame, context: ContextTypes.DEFAULT_TYPE
):
    df = dataframe.copy().reset_index(drop=True)
    headers = list(df.columns)

    date_col_name: Optional[str] = None
    for h in headers:
        if "дата выезда" in str(h).lower():
            date_col_name = h
            break
    if date_col_name:
        try:
            df[date_col_name] = pd.to_datetime(
                df[date_col_name], errors="coerce", dayfirst=True
            )
        except Exception:
            pass

    settings = get_schedule_state()
    version = get_schedule_version(settings)
    approvals = get_schedule_approvals(version)

    bio = BytesIO()
    with pd.ExcelWriter(bio, engine="openpyxl") as writer:
        df.to_excel(
            writer,
            sheet_name="График выездов",
            index=False,
            startrow=2,
            header=False,
        )

        wb = writer.book
        ws = writer.sheets["График выездов"]

        for col_num, value in enumerate(headers, 1):
            cell = ws.cell(row=2, column=col_num, value=value)
            cell.fill = HEADER_FILL
            cell.font = HEADER_FONT
            cell.alignment = Alignment(horizontal="center", vertical="center")

        for column in ws.columns:
            max_length = 0
            col_letter = column[0].column_letter
            for cell in column:
                try:
                    if cell.value is not None and len(str(cell.value)) > max_length:
                        max_length = len(str(cell.value))
                except Exception:
                    pass
            ws.column_dimensions[col_letter].width = min(max_length + 4, 50)

        ws.freeze_panes = ws["A3"]

        last_col_letter = ws.cell(row=2, column=len(headers)).column_letter
        ws.auto_filter.ref = f"A2:{last_col_letter}{len(df) + 2}"

        for row in ws[f"A3:{last_col_letter}{len(df) + 2}"]:
            for cell in row:
                cell.border = BORDER

        LIGHT_FILL = PatternFill(
            start_color="F0F0F0", end_color="F0F0F0", fill_type="solid"
        )
        for idx, row in enumerate(
            ws.iter_rows(min_row=3, max_row=len(df) + 2), start=3
        ):
            if idx % 2 == 0:
                for cell in row:
                    cell.fill = LIGHT_FILL

        tab = Table(
            displayName="ScheduleTable",
            ref=f"A2:{last_col_letter}{len(df) + 2}",
        )
        tab.tableStyleInfo = TableStyleInfo(
            name="TableStyleMedium9",
            showFirstColumn=False,
            showLastColumn=False,
            showRowStripes=True,
            showColumnStripes=False,
        )
        ws.add_table(tab)

        date_idx = None
        onzs_idx = None
        dev_idx = None
        obj_idx = None

        for i, h in enumerate(headers, start=1):
            h_low = str(h).lower()
            if date_idx is None and "дата выезда" in h_low:
                date_idx = i
            if onzs_idx is None and "онзс" in h_low:
                onzs_idx = i
            if dev_idx is None and "наименование застройщика" in h_low:
                dev_idx = i
            if obj_idx is None and "наименование объекта" in h_low:
                obj_idx = i

        for row_idx in range(3, len(df) + 3):
            if date_idx:
                cell = ws.cell(row=row_idx, column=date_idx)
                cell.number_format = "DD.MM.YYYY"
            if onzs_idx:
                cell = ws.cell(row=row_idx, column=onzs_idx)
                cell.alignment = Alignment(
                    horizontal="center", vertical="center", wrap_text=False
                )
            if dev_idx:
                cell = ws.cell(row=row_idx, column=dev_idx)
                cell.alignment = Alignment(
                    horizontal="left", vertical="center", wrap_text=True
                )
            if obj_idx:
                cell = ws.cell(row=row_idx, column=obj_idx)
                cell.alignment = Alignment(
                    horizontal="left", vertical="center", wrap_text=True
                )

        if approvals:
            last_data_row = len(df) + 2
            summary_start = last_data_row + 2

            header = build_schedule_header(version, approvals)
            ws.merge_cells(f"A{summary_start}:{last_col_letter}{summary_start}")
            cell_header = ws[f"A{summary_start}"]
            cell_header.value = header
            cell_header.font = Font(bold=True, size=12, color="FFFFFF")
            cell_header.fill = PatternFill(
                start_color="4F81BD", end_color="4F81BD", fill_type="solid"
            )
            cell_header.alignment = Alignment(horizontal="center", vertical="center")

            sub_row = summary_start + 1
            ws.merge_cells(f"A{sub_row}:{last_col_letter}{sub_row}")
            cell_sub = ws[f"A{sub_row}"]
            cell_sub.value = "Согласовано всеми:"
            cell_sub.font = Font(bold=True, size=11)
            cell_sub.alignment = Alignment(horizontal="left", vertical="center")

            row_ptr = sub_row + 1
            approved_rows = [r for r in approvals if r["status"] == "approved"]
            others = [r for r in approvals if r["status"] != "approved"]

            list_fill = PatternFill(
                start_color="D9E1F2", end_color="D9E1F2", fill_type="solid"
            )

            for r in approved_rows:
                line = f"• {r['approver']} — {_format_dt(r['decided_at'])} ✅"
                ws.merge_cells(f"A{row_ptr}:{last_col_letter}{row_ptr}")
                cell = ws[f"A{row_ptr}"]
                cell.value = line
                cell.fill = list_fill
                cell.font = Font(size=11)
                cell.alignment = Alignment(horizontal="left", vertical="center")
                for col_idx in range(1, len(headers) + 1):
                    ws.cell(row=row_ptr, column=col_idx).border = BORDER
                row_ptr += 1

            if others:
                ws.merge_cells(f"A{row_ptr}:{last_col_letter}{row_ptr}")
                cell_pending = ws[f"A{row_ptr}"]
                cell_pending.value = "⚠ Есть несогласованные/на доработке."
                cell_pending.font = Font(italic=True, color="C00000")
                cell_pending.alignment = Alignment(
                    horizontal="left", vertical="center"
                )
                for col_idx in range(1, len(headers) + 1):
                    ws.cell(row=row_ptr, column=col_idx).border = BORDER

    bio.seek(0)
    filename = f"График_выездов_СОТ_{date.today().strftime('%d.%m.%Y')}.xlsx"

    await context.bot.send_document(
        chat_id=chat_id,
        document=InputFile(bio, filename=filename),
        caption="График выездов отдела СОТ",
    )


# -------------------------------------------------
# Текст графика
# -------------------------------------------------
def _format_dt(iso_str: Optional[str]) -> str:
    if not iso_str:
        return ""
    try:
        dt = datetime.fromisoformat(iso_str)
        return dt.strftime("%d.%m.%Y %H:%M")
    except Exception:
        return iso_str


def _compute_schedule_dates(
    approvals: List[sqlite3.Row],
) -> (Optional[date], Optional[date]):
    dates: List[date] = []
    for r in approvals:
        if r["status"] == "approved" and r["decided_at"]:
            try:
                dt = datetime.fromisoformat(r["decided_at"])
                dates.append(dt.date())
            except Exception:
                pass
    if not dates:
        return None, None
    base = max(dates)
    d_from = base
    d_to = base + timedelta(days=4)
    return d_from, d_to


def build_schedule_header(version: int, approvals: List[sqlite3.Row]) -> str:
    d_from, d_to = _compute_schedule_dates(approvals)
    if not d_from or not d_to:
        return f"📅 График выездов (версия {version})"
    return f"📅 График выездов с {d_from:%d.%m.%Y} по {d_to:%d.%m.%Y} г"


def write_schedule_summary_to_sheet(version: int, approvals: List[sqlite3.Row]) -> None:
    service = get_sheets_service()
    if service is None:
        log.error(
            "Google Sheets сервис недоступен – не могу записать итог согласования в 'График'."
        )
    else:
        sheet_name = "График"
        header = build_schedule_header(version, approvals)
        rows = [
            [""],
            [header],
            ["Согласовано всеми:"],
        ]
        for r in approvals:
            rows.append(
                [f"{r['approver']} — {_format_dt(r['decided_at'])} ✅"]
            )

        body = {"values": rows}

        try:
            service.spreadsheets().values().append(
                spreadsheetId=GSHEETS_SPREADSHEET_ID,
                range=f"'{sheet_name}'!A1",
                valueInputOption="USER_ENTERED",
                insertDataOption="INSERT_ROWS",
                body=body,
            ).execute()
            log.info(
                "Итог согласования версии %s дописан в лист '%s'.",
                version,
                sheet_name,
            )
        except Exception as e:
            log.error(
                "Ошибка записи итога согласования в лист '%s': %s", sheet_name, e
            )


def build_schedule_text(is_admin_flag: bool, settings: dict) -> str:
    version = get_schedule_version(settings)
    approvals = get_schedule_approvals(version)
    approvers = get_current_approvers(settings)

    header = build_schedule_header(version, approvals)
    lines = [header, ""]

    if not approvers:
        lines.append("Согласующие не назначены.")
        return "\n".join(lines)

    pending: List[str] = []
    approved: List[sqlite3.Row] = []
    rework: List[sqlite3.Row] = []

    by_approver = {r["approver"]: r for r in approvals}

    for a in approvers:
        r = by_approver.get(a)
        if not r or r["status"] == "pending":
            pending.append(a)
        elif r["status"] == "approved":
            approved.append(r)
        elif r["status"] == "rework":
            rework.append(r)

    if rework:
        lines.append("Отправлено на доработку:")
        for r in rework:
            lines.append(
                f"• {r['approver']} — {_format_dt(r['decided_at'])} (Комментарий: {r['comment'] or 'нет'})"
            )
    elif pending:
        lines.append("На согласовании у:")
        for a in pending:
            lines.append(
                f"• {a} — запрошено {_format_dt(by_approver[a]['requested_at'])}"
            )
        if approved:
            lines.append("")
            lines.append("Уже согласовали:")
            for r in approved:
                lines.append(f"• {r['approver']} — {_format_dt(r['decided_at'])} ✅")
    else:
        lines.append("Согласовано всеми:")
        for r in approved:
            lines.append(f"• {r['approver']} — {_format_dt(r['decided_at'])} ✅")

    return "\n".join(lines)


# -------------------------------------------------
# Замечания: НЕ УСТРАНЕНЫ
# -------------------------------------------------
def build_remarks_not_done_text(df: pd.DataFrame) -> str:
    COLS = {
        "case": "I",
        "pb": "Q",
        "pb_zk": "R",
        "ar": "X",
        "eom": "AD",
    }

    TITLES = {
        "pb": "Отметка об устранении замечаний ПБ да/нет",
        "pb_zk": "Отметка об устранении замечаний ПБ в ЗК КНД да/нет",
        "ar": "Отметка об устранении нарушений АР, ММГН, АГО да/нет",
        "eom": "Отметка об устранении нарушений ЭОМ да/нет",
    }

    idx_case = excel_col_to_index(COLS["case"])
    idx_pb = excel_col_to_index(COLS["pb"])
    idx_pb_zk = excel_col_to_index(COLS["pb_zk"])
    idx_ar = excel_col_to_index(COLS["ar"])
    idx_eom = excel_col_to_index(COLS["eom"])

    def is_net(val):
        if val is None:
            return False
        text = str(val).lower().replace("\n", " ").strip()
        if not text or text in {"-", "н/д"}:
            return False
        return text.startswith("нет")

    grouped: Dict[str, Dict[str, Any]] = {}

    for _, row in df.iterrows():
        case = str(row.iloc[idx_case]).strip() if idx_case < len(row) else ""
        if not case:
            continue

        sheet_src = ""
        try:
            sheet_src = str(row.get("_remarks_sheet", "")).strip()
        except Exception:
            sheet_src = ""

        flags = {
            "pb": is_net(row.iloc[idx_pb]) if idx_pb < len(row) else False,
            "pb_zk": is_net(row.iloc[idx_pb_zk]) if idx_pb_zk < len(row) else False,
            "ar": is_net(row.iloc[idx_ar]) if idx_ar < len(row) else False,
            "eom": is_net(row.iloc[idx_eom]) if idx_eom < len(row) else False,
        }

        if not any(flags.values()):
            continue

        if case not in grouped:
            grouped[case] = {"pb": set(), "ar": set(), "eom": set(), "sheets": set()}

        if sheet_src:
            grouped[case]["sheets"].add(sheet_src)

        if flags["pb"]:
            grouped[case]["pb"].add(TITLES["pb"])
        if flags["pb_zk"]:
            grouped[case]["pb"].add(TITLES["pb_zk"])
        if flags["ar"]:
            grouped[case]["ar"].add(TITLES["ar"])
        if flags["eom"]:
            grouped[case]["eom"].add(TITLES["eom"])

    if not grouped:
        return "Во всех строках нет статусов «нет»."

    sheets_present: List[str] = []
    try:
        sheets_present = sorted(
            {
                str(x).strip()
                for x in df.get("_remarks_sheet", pd.Series([])).dropna().unique().tolist()
                if str(x).strip()
            }
        )
    except Exception:
        sheets_present = []

    sheets_line = " / ".join(sheets_present) if sheets_present else "—"

    lines = [
        "Строки со статусом «НЕ УСТРАНЕНЫ (нет)»",
        "",
        "Листы: " + sheets_line,
        "",
    ]

    for case, blocks in grouped.items():
        parts = []
        if blocks["pb"]:
            parts.append("Пожарная безопасность: " + ", ".join(b + " - нет" for b in blocks["pb"]))
        if blocks["ar"]:
            parts.append("Архитектура, ММГН, АГО: " + ", ".join(b + " - нет" for b in blocks["ar"]))
        if blocks["eom"]:
            parts.append("Электроснабжение: " + ", ".join(b + " - нет" for b in blocks["eom"]))

        src = ""
        if blocks.get("sheets"):
            src = " (" + " / ".join(sorted(blocks["sheets"])) + ")"

        lines.append(f"• {case}{src} — " + "; ".join(parts))

    return "\n".join(lines)

def build_remarks_not_done_by_onzs(df: pd.DataFrame, onzs_value: str) -> str:
    onzs_idx = get_col_index_by_header(df, "онзс", "D")
    if onzs_idx is None:
        return "Не удалось определить столбец ОНзС в файле замечаний."

    COLS = {
        "case": "I",
        "pb": "Q",
        "pb_zk": "R",
        "ar": "X",
        "eom": "AD",
    }

    TITLES = {
        "pb": "Отметка об устранении замечаний ПБ да/нет",
        "pb_zk": "Отметка об устранении замечаний ПБ в ЗК КНД да/нет",
        "ar": "Отметка об устранении нарушений АР, ММГН, АГО да/нет",
        "eom": "Отметка об устранении нарушений ЭОМ да/нет",
    }

    idx_case = excel_col_to_index(COLS["case"])
    idx_pb = excel_col_to_index(COLS["pb"])
    idx_pb_zk = excel_col_to_index(COLS["pb_zk"])
    idx_ar = excel_col_to_index(COLS["ar"])
    idx_eom = excel_col_to_index(COLS["eom"])

    def is_net(val):
        if val is None:
            return False
        text = str(val).lower().replace("\n", " ").strip()
        if not text or text in {"-", "н/д"}:
            return False
        return text.startswith("нет")

    grouped: Dict[str, Dict[str, Any]] = {}

    num_str = normalize_onzs_value(onzs_value)

    for _, row in df.iterrows():
        try:
            val_raw = row.iloc[onzs_idx]
        except Exception:
            val_raw = None

        val_norm = normalize_onzs_value(val_raw)
        if val_norm != num_str:
            continue

        case = ""
        try:
            case = str(row.iloc[idx_case]).strip()
        except Exception:
            pass

        if not case:
            continue

        sheet_src = ""
        try:
            sheet_src = str(row.get("_remarks_sheet", "")).strip()
        except Exception:
            sheet_src = ""

        flags = {
            "pb": is_net(row.iloc[idx_pb]) if idx_pb < len(row) else False,
            "pb_zk": is_net(row.iloc[idx_pb_zk]) if idx_pb_zk < len(row) else False,
            "ar": is_net(row.iloc[idx_ar]) if idx_ar < len(row) else False,
            "eom": is_net(row.iloc[idx_eom]) if idx_eom < len(row) else False,
        }

        if not any(flags.values()):
            continue

        if case not in grouped:
            grouped[case] = {"pb": set(), "ar": set(), "eom": set(), "sheets": set()}

        if sheet_src:
            grouped[case]["sheets"].add(sheet_src)

        if flags["pb"]:
            grouped[case]["pb"].add(TITLES["pb"])
        if flags["pb_zk"]:
            grouped[case]["pb"].add(TITLES["pb_zk"])
        if flags["ar"]:
            grouped[case]["ar"].add(TITLES["ar"])
        if flags["eom"]:
            grouped[case]["eom"].add(TITLES["eom"])

    sheets_present: List[str] = []
    try:
        sheets_present = sorted(
            {
                str(x).strip()
                for x in df.get("_remarks_sheet", pd.Series([])).dropna().unique().tolist()
                if str(x).strip()
            }
        )
    except Exception:
        sheets_present = []

    sheets_line = " / ".join(sheets_present) if sheets_present else "—"

    if not grouped:
        return (
            f"По ОНзС {onzs_value} нет строк со статусом «нет».\n"
            f"Листы: {sheets_line}"
        )

    lines = [
        f"Строки со статусом «НЕ УСТРАНЕНЫ (нет)» по ОНзС {onzs_value}",
        "",
        "Листы: " + sheets_line,
        "",
    ]

    for case, blocks in grouped.items():
        parts = []
        if blocks["pb"]:
            parts.append("Пожарная безопасность: " + ", ".join(b + " - нет" for b in blocks["pb"]))
        if blocks["ar"]:
            parts.append("Архитектура, ММГН, АГО: " + ", ".join(b + " - нет" for b in blocks["ar"]))
        if blocks["eom"]:
            parts.append("Электроснабжение: " + ", ".join(b + " - нет" for b in blocks["eom"]))

        src = ""
        if blocks.get("sheets"):
            src = " (" + " / ".join(sorted(blocks["sheets"])) + ")"

        lines.append(f"• {case}{src} — " + "; ".join(parts))

    return "\n".join(lines)

def build_case_cards_text(df: pd.DataFrame, case_no: str) -> str:
    case_no = case_no.strip()
    if not case_no:
        return "Номер дела не указан."

    target = normalize_case_number(case_no)

    idx_case = get_case_col_index(df)
    if idx_case is None:
        return (
            "Не удалось определить столбец «Номер дела (I)» в файле замечаний. "
            "Проверьте структуру листа."
        )

    idx_date = get_col_index_by_header(df, "дата выезда", "B")
    idx_onzs = get_col_index_by_header(df, "онзс", "D")
    idx_dev = get_col_index_by_header(df, "наименование застройщика", "F")
    idx_obj = get_col_index_by_header(df, "наименование объекта", "G")
    idx_addr = get_col_index_by_header(df, "строительный адрес", "H")

    idx_pb = excel_col_to_index("Q")
    idx_pb_zk = excel_col_to_index("R")
    idx_ar = excel_col_to_index("X")
    idx_eom = excel_col_to_index("AD")

    mask: List[bool] = []
    for _, row in df.iterrows():
        try:
            val_raw = row.iloc[idx_case]
        except Exception:
            val_raw = None
        val_norm = normalize_case_number(val_raw)
        mask.append(val_norm == target)

    if not any(mask):
        return (
            f"По номеру дела {case_no} ничего не найдено.\n"
            f"Листы: " + " / ".join([s for s in get_remarks_sheet_candidates()])
        )

    df_sel = df[mask]


    # Листы (источники) в выборке
    sheets_sel: List[str] = []
    try:
        sheets_sel = sorted(
            {
                str(x).strip()
                for x in df_sel.get("_remarks_sheet", pd.Series([])).dropna().unique().tolist()
                if str(x).strip()
            }
        )
    except Exception:
        sheets_sel = []
    sheets_line = " / ".join(sheets_sel) if sheets_sel else "—"

    lines: List[str] = [
        f"Результаты поиска по номеру дела: {case_no}",
        "",
        "Листы: " + sheets_line,
        "",
    ]
    for _, row in df_sel.iterrows():

        def safe(idx: Optional[int]) -> str:
            if idx is None:
                return ""
            try:
                return str(row.iloc[idx]).strip()
            except Exception:
                return ""

        date_raw = safe(idx_date)
        date_fmt = date_raw
        try:
            if date_raw:
                dt = pd.to_datetime(date_raw, dayfirst=True, errors="ignore")
                if isinstance(dt, (datetime, pd.Timestamp)):
                    date_fmt = dt.strftime("%d.%m.%Y")
        except Exception:
            pass

        onzs_val = safe(idx_onzs)
        dev_val = safe(idx_dev)
        obj_val = safe(idx_obj)
        addr_val = safe(idx_addr)

        def safe_status(idx: int) -> str:
            try:
                if idx < len(row):
                    return str(row.iloc[idx]).strip()
            except Exception:
                pass
            return ""

        pb_val = safe_status(idx_pb)
        pb_zk_val = safe_status(idx_pb_zk)
        ar_val = safe_status(idx_ar)
        eom_val = safe_status(idx_eom)

        lines.append(f"Номер дела: {case_no}")
        try:
            sheet_src = str(row.get("_remarks_sheet", "")).strip()
        except Exception:
            sheet_src = ""
        if sheet_src:
            lines.append(f"Лист: {sheet_src}")
        if date_fmt:
            lines.append(f"Дата выезда: {date_fmt}")
        if onzs_val:
            lines.append(f"ОНзС: {onzs_val}")
        if dev_val:
            lines.append(f"Застройщик: {dev_val}")
        if obj_val:
            lines.append(f"Объект: {obj_val}")
        if addr_val:
            lines.append(f"Адрес: {addr_val}")

        lines.append("")
        lines.append(f"ПБ: {pb_val or '-'}")
        lines.append(f"ПБ ЗК: {pb_zk_val or '-'}")
        lines.append(f"АР/ММГН/АГО: {ar_val or '-'}")
        lines.append(f"ЭОМ: {eom_val or '-'}")
        lines.append("")
        lines.append("────────────")
        lines.append("")

    return "\n".join(lines)


# -------------------------------------------------
# Отправка длинного текста
# -------------------------------------------------
async def send_long_text(chat, text: str, chunk_size=3500):
    lines = text.split("\n")
    buf = ""

    for line in lines:
        if len(buf) + len(line) + 1 > chunk_size:
            await chat.send_message(buf)
            buf = line
        else:
            buf = buf + "\n" + line if buf else line

    if buf:
        await chat.send_message(buf)


# -------------------------------------------------
# Лист замечаний (поддержка нескольких лет)
# -------------------------------------------------

def get_remarks_sheet_candidates() -> List[str]:
    """
    Возвращает список кандидатов листов замечаний.
    По умолчанию берём текущий год и предыдущий год (например, 2026 и 2025),
    чтобы замечания подтягивались сразу за два года.
    """
    y = local_now().year
    years = [y, y - 1]
    return [f"ПБ, АР,ММГН, АГО ({yy})" for yy in years]


def get_remarks_df_current() -> Optional[pd.DataFrame]:
    """
    Читает лист(ы) замечаний из основной таблицы GSHEETS_SPREADSHEET_ID.

    Важно:
    - бот подтягивает замечания не только за текущий год, но и за предыдущий (например, 2026 + 2025);
    - строки помечаются служебной колонкой _remarks_sheet, чтобы в ответах можно было показать источник.
    """
    url = build_export_url(GSHEETS_SPREADSHEET_ID)

    try:
        resp = requests.get(url, timeout=30)
        resp.raise_for_status()
        xls = pd.ExcelFile(BytesIO(resp.content))

        frames: List[pd.DataFrame] = []
        for sheet in get_remarks_sheet_candidates():
            if sheet not in xls.sheet_names:
                continue
            df = pd.read_excel(xls, sheet_name=sheet)
            df = df.dropna(how="all")
            if df.empty:
                continue
            df["_remarks_sheet"] = sheet
            frames.append(df)

        if not frames:
            fallback_sheet = get_current_remarks_sheet_name()
            if fallback_sheet not in xls.sheet_names:
                log.error("В файле нет листов замечаний: %s", ", ".join(get_remarks_sheet_candidates()))
                return None
            df = pd.read_excel(xls, sheet_name=fallback_sheet)
            df["_remarks_sheet"] = fallback_sheet
            return df

        return pd.concat(frames, ignore_index=True)

    except Exception as e:
        log.error("Ошибка чтения листа(ов) замечаний: %s", e)
        return None


# -------------------------------------------------
# Итоговые проверки: чтение, фильтр, текст, Excel
# -------------------------------------------------
def refresh_final_checks_local_file() -> bool:
    """
    Обновляет локальный файл итоговых проверок:
    - удаляет старый файл (если есть);
    - скачивает актуальную версию из Google Sheets по FINAL_CHECKS_SPREADSHEET_ID.
    """
    sheet_id = FINAL_CHECKS_SPREADSHEET_ID
    if not sheet_id:
        log.error("FINAL_CHECKS_SPREADSHEET_ID не задан.")
        return False

    url = build_export_url(sheet_id)
    path = FINAL_CHECKS_LOCAL_PATH

    # удаляем старый файл, если есть
    try:
        if os.path.exists(path):
            os.remove(path)
    except Exception as e:
        log.warning(
            "Не удалось удалить старый файл итоговых проверок %s: %s",
            path,
            e,
        )

    try:
        resp = requests.get(url, timeout=30)
        resp.raise_for_status()
    except Exception as e:
        log.error("Ошибка скачивания Excel (итоговые проверки): %s", e)
        return False

    try:
        with open(path, "wb") as f:
            f.write(resp.content)
        log.info("Файл итоговых проверок сохранён локально: %s", path)
        return True
    except Exception as e:
        log.error(
            "Ошибка записи локального файла итоговых проверок %s: %s",
            path,
            e,
        )
        return False


def get_final_checks_df() -> Optional[pd.DataFrame]:
    """
    Читает локальный файл итоговых проверок FINAL_CHECKS_LOCAL_PATH,
    который обновляется при входе в раздел «Итоговые проверки».
    Собирает данные со всех листов книги и склеивает их в один DataFrame.
    """
    path = FINAL_CHECKS_LOCAL_PATH
    if not path:
        log.error("FINAL_CHECKS_LOCAL_PATH не задан.")
        return None

    if not os.path.exists(path):
        log.error("Локальный файл итоговых проверок не найден: %s", path)
        return None

    try:
        xls = pd.ExcelFile(path)
        if not xls.sheet_names:
            log.error("Файл итоговых проверок пуст (нет листов).")
            return None

        frames: List[pd.DataFrame] = []
        for sheet_name in xls.sheet_names:
            try:
                df_sheet = pd.read_excel(xls, sheet_name=sheet_name)
                df_sheet = df_sheet.dropna(how="all")
                if not df_sheet.empty:
                    frames.append(df_sheet)
            except Exception as e:
                log.warning(
                    "Ошибка чтения листа '%s' итоговых проверок: %s",
                    sheet_name,
                    e,
                )

        if not frames:
            log.error("Во всех листах итоговых проверок нет данных.")
            return None

        df = pd.concat(frames, ignore_index=True)
        df = df.reset_index(drop=True)
        return df
    except Exception as e:
        log.error("Ошибка чтения локального файла итоговых проверок: %s", e)
        return None




def _parse_final_date(val) -> Optional[date]:
    """
    Универсальное преобразование значения из столбцов O/P в дату.
    Поддерживает:
    - datetime / Timestamp;
    - текстовые даты ("22.12.2025", "22.12.25", "22.12", "22.12.2025 г.");
    - числовые Excel‑серийные даты (целое/float).
    Всегда возвращает date или None.
    """
    # 1. Пустые / отсутствующие значения
    if val is None or (isinstance(val, float) and pd.isna(val)) or (not isinstance(val, (int, float, datetime, pd.Timestamp)) and pd.isna(val)):
        return None

    # 2. Уже datetime / Timestamp
    if isinstance(val, (datetime, pd.Timestamp)):
        try:
            year = val.year
        except Exception:
            return None
        if year < 1900 or year > 2100:
            return None
        return val.date()

    # 3. Чисто числовые значения — Excel serial (float/int)
    if isinstance(val, (int, float)):
        serial = float(val)
        # отсекаем заведомо мусорные значения (как в предупреждениях openpyxl)
        if 20000 <= serial <= 80000:  # ~1945–2120 гг.
            excel_epoch = date(1899, 12, 30)
            return excel_epoch + timedelta(days=int(serial))
        else:
            return None

    # 4. Всё остальное — пробуем парсить как строку
    s = str(val).strip()
    if not s:
        return None

    # Нормализация строки
    s_norm = (
        s.replace("г.", "")
         .replace("г", "")
         .replace("\xa0", "")
         .replace(" ", "")
         .replace("–", "-")
         .replace("—", "-")
    )

    # Если осталась только 1–2 цифры — это не дата
    if re.fullmatch(r"\d{1,2}$", s_norm):
        return None

    # 5. Явные текстовые даты
    for fmt in ("%d.%m.%Y", "%d.%m.%y", "%d.%m"):
        try:
            dt_obj = datetime.strptime(s_norm, fmt)
            if fmt == "%d.%m":
                dt_obj = dt_obj.replace(year=date.today().year)
            if dt_obj.year < 1900 or dt_obj.year > 2100:
                return None
            return dt_obj.date()
        except ValueError:
            continue

    # 6. Строковый Excel‑серийный номер
    if re.fullmatch(r"\d{4,6}", s_norm):
        try:
            serial = float(s_norm)
            if 20000 <= serial <= 80000:
                excel_epoch = date(1899, 12, 30)
                return excel_epoch + timedelta(days=int(serial))
        except Exception:
            pass

    return None


# -------------------------------------------------
# 🚨 Красные лампочки — загрузка Excel и BI-панель
# -------------------------------------------------

REDLAMPS_TOLERANCE_DAYS_DEFAULT = 5

REDLAMPS_LOW_RISK_KEYWORDS = [
    "очистные",
    "котельн",
    "дорог",
    "мост",
    "тп",
    "линейн",
    "взу",
    "канализац",
    "лэп",
    "инженерн",
]

def _rl_has_text(val: Any) -> bool:
    if val is None:
        return False
    try:
        if isinstance(val, float) and pd.isna(val):
            return False
    except Exception:
        pass
    s = str(val).strip()
    if not s or s.lower() == "nan":
        return False
    return True


def _rl_cell_filled(val: Any) -> bool:
    """Строгая проверка заполненности ячейки (для O/T)."""
    if val is None:
        return False
    try:
        if pd.isna(val):
            return False
    except Exception:
        pass
    s = str(val).strip()
    if not s:
        return False
    bad = {"nan", "none", "null", "<na>", "nat", "-", "—"}
    if s.lower() in bad:
        return False
    return True

def _rl_contains(val: Any, needle: str) -> bool:
    if not _rl_has_text(val):
        return False
    return needle.lower() in str(val).lower()

def _rl_contains_any_keywords(val: Any, keywords: List[str]) -> bool:
    if not _rl_has_text(val):
        return False
    s = str(val).lower()
    return any(k in s for k in keywords)

def redlamps_menu_inline() -> InlineKeyboardMarkup:
    return InlineKeyboardMarkup(
        [
            [InlineKeyboardButton("📤 Загрузить файл", callback_data="redlamps_upload")],
            [InlineKeyboardButton("📅 Выбрать период (K–L)", callback_data="redlamps_period")],
            [InlineKeyboardButton("📊 Сформировать BI-панель", callback_data="redlamps_build")],
            [InlineKeyboardButton("🗑 Сбросить", callback_data="redlamps_reset")],
        ]
    )

def _redlamps_process_bytes(
    xlsx_bytes: bytes,
    date_from: date,
    date_to: date,
    tolerance_days: int = REDLAMPS_TOLERANCE_DAYS_DEFAULT,
) -> str:
    # читаем без заголовков, чтобы работать по буквам колонок как в Excel
    df = pd.read_excel(BytesIO(xlsx_bytes), header=None)
    if df is None or df.empty:
        return "Файл пустой или не удалось прочитать Excel."

    # Индексы колонок (0-based) по буквам Excel
    IDX_B = excel_col_to_index("B")
    IDX_E = excel_col_to_index("E")
    IDX_I = excel_col_to_index("I")
    IDX_K = excel_col_to_index("K")
    IDX_L = excel_col_to_index("L")
    IDX_O = excel_col_to_index("O")
    IDX_T = excel_col_to_index("T")
    IDX_AA = excel_col_to_index("AA")

    # защита от файлов, где меньше колонок
    max_idx = max(IDX_B, IDX_E, IDX_I, IDX_K, IDX_L, IDX_O, IDX_T, IDX_AA)
    if len(df.columns) <= max_idx:
        return (
            "В загруженном файле недостаточно колонок для обработки.\n"
            "Нужно, чтобы были колонки как минимум до AA (включая B, E, I, K, L, O, T, AA)."
        )

    # 1) Удаляем строки по AA: «Завершение строительства»
    aa = df.iloc[:, IDX_AA]
    mask_aa = aa.apply(lambda v: _rl_contains(v, "Завершение строительства"))
    df = df[~mask_aa]

    # 2) Удаляем строки по E (низкий риск)
    e = df.iloc[:, IDX_E]
    mask_low_risk = e.apply(lambda v: _rl_contains_any_keywords(v, REDLAMPS_LOW_RISK_KEYWORDS))
    df = df[~mask_low_risk]

    # 3) Удаляем все строки по I кроме «Программа проверок»
    i_col = df.iloc[:, IDX_I]
    mask_program = i_col.apply(lambda v: _rl_contains(v, "Программа проверок"))
    df = df[mask_program]

    if df.empty:
        return (
            "После обязательных фильтров (AA/низкий риск E/только «Программа проверок» в I) "
            "не осталось строк."
        )

    # 4) Фильтр по периоду K/L с допуском ±N дней
    k_raw = df.iloc[:, IDX_K]
    l_raw = df.iloc[:, IDX_L]

    k_dates = k_raw.apply(_parse_final_date)
    l_dates = l_raw.apply(_parse_final_date)

    start = date_from - timedelta(days=tolerance_days)
    end = date_to + timedelta(days=tolerance_days)

    def _in_range(d: Optional[date]) -> bool:
        if not d:
            return False
        return start <= d <= end

    mask_period = k_dates.apply(_in_range) | l_dates.apply(_in_range)
    df = df[mask_period]

    if df.empty:
        return (
            f"По периоду {date_from:%d.%m.%Y} — {date_to:%d.%m.%Y} (±{tolerance_days} дней) "
            "ничего не найдено."
        )

    # 5) Группировка по номеру дела (B) + подсчёты O/T
    stats: Dict[str, Dict[str, int]] = {}
    for _, row in df.iterrows():
        case_norm = normalize_case_number(row.iloc[IDX_B])
        if not case_norm:
            continue

        acts_inc = 1 if _rl_cell_filled(row.iloc[IDX_O]) else 0
        prot_inc = 1 if _rl_cell_filled(row.iloc[IDX_T]) else 0

        if case_norm not in stats:
            stats[case_norm] = {"acts": 0, "protocols": 0}

        stats[case_norm]["acts"] += acts_inc
        stats[case_norm]["protocols"] += prot_inc

    if not stats:
        return "Не удалось собрать номера дел (колонка B) после фильтров."

    # 6) BI-панель
    total_cases = len(stats)
    total_acts = sum(v["acts"] for v in stats.values())
    total_prot = sum(v["protocols"] for v in stats.values())
    bad = 0

    lines: List[str] = []
    lines.append("🚨 Красные лампочки — BI-панель")
    lines.append(f"Период: {date_from:%d.%m.%Y} — {date_to:%d.%m.%Y} (±{tolerance_days} дней)")
    lines.append(f"Дел: {total_cases} | Актов: {total_acts} | Протоколов: {total_prot}")
    lines.append("")

    for case_no in sorted(stats.keys()):
        acts = stats[case_no]["acts"]
        prots = stats[case_no]["protocols"]

        # правило: если актов >= 4, протоколов должно быть минимум floor(acts/2)
        if acts >= 4:
            need = acts // 2
            if prots < need:
                bad += 1
                lines.append(f"❗ {case_no} — Акты: {acts}; Протоколы: {prots} (нужно ≥ {need})")
            else:
                lines.append(f"✅ {case_no} — Акты: {acts}; Протоколы: {prots}")
        else:
            lines.append(f"✅ {case_no} — Акты: {acts}; Протоколы: {prots}")

    lines.insert(4, f"Проблемных дел: {bad}")

    return "\n".join(lines)
def filter_final_checks_df(
    df: pd.DataFrame,
    start_date: Optional[date] = None,
    end_date: Optional[date] = None,
    case_no: Optional[str] = None,
    basis: str = "any",  # "start" -> только O, "end" -> только P, "any" -> O или P
) -> pd.DataFrame:
    """
    Универсальный фильтр итоговых проверок:
    - по периоду (O / P в зависимости от basis);
    - по номеру дела.
    Работает в связке с кнопками:
      • За неделю / За месяц (basis = "start" или "end");
      • Выбрать период;
      • По номеру дела.
    """
    if df is None or df.empty:
        return df.iloc[0:0].copy()

    # Индексы колонок в итоговой таблице:
    # B — номер дела, O — дата начала, P — дата окончания
    idx_case = excel_col_to_index("B")
    idx_start = excel_col_to_index("O")
    idx_end = excel_col_to_index("P")

    basis = (basis or "any").lower()
    if basis not in ("start", "end", "any"):
        basis = "any"

    result = df.copy()

    # ---------- Фильтр по номеру дела ----------
    if case_no:
        case_filter_norm = normalize_case_number(case_no)
        if not case_filter_norm:
            return result.iloc[0:0].copy()

        try:
            ser_case = result.iloc[:, idx_case]
        except Exception:
            # если вдруг нет колонки B — возвращаем пустой df
            return result.iloc[0:0].copy()

        def _norm(v):
            return normalize_case_number(v)

        mask_case = ser_case.apply(lambda v: _norm(v) == case_filter_norm)
        result = result[mask_case]

        if result.empty:
            return result

    # ---------- Фильтр по датам O/P ----------
    if start_date or end_date:
        # берём "сырые" значения из O и P
        try:
            ser_start_raw = result.iloc[:, idx_start]
        except Exception:
            ser_start_raw = pd.Series([None] * len(result), index=result.index)

        try:
            ser_end_raw = result.iloc[:, idx_end]
        except Exception:
            ser_end_raw = pd.Series([None] * len(result), index=result.index)

        # приводим каждое значение к date (или None)
        ser_start = ser_start_raw.apply(_parse_final_date)
        ser_end = ser_end_raw.apply(_parse_final_date)

        # выбираем базовую дату для фильтра
        if basis == "start":
            base = ser_start
        elif basis == "end":
            base = ser_end
        else:  # "any" — сначала O, если пусто, берём P
            base = ser_start.where(ser_start.notna(), ser_end)

        # переводим в Timestamp, чтобы можно было сравнивать диапазон
        base_dt = pd.to_datetime(base, errors="coerce")

        mask = pd.Series(True, index=result.index)
        if start_date:
            mask &= base_dt >= pd.to_datetime(start_date)
        if end_date:
            mask &= base_dt <= pd.to_datetime(end_date)

        result = result[mask]

    return result.reset_index(drop=True)



def compute_auto_period_for_final(df: pd.DataFrame, basis: str, mode: str) -> Optional[tuple[date, date]]:
    """Определяет автоматически период для итоговых проверок.

    basis:
        'start' — использовать столбец O (дата начала итоговой проверки)
        'end'   — использовать столбец P (дата окончания итоговой проверки)

    mode:
        'week'  — последние 7 дней от максимальной даты
        'month' — последние 30 дней от максимальной даты
    """
    if df is None or df.empty:
        return None

    basis = (basis or "start").lower()
    if basis not in ("start", "end"):
        basis = "start"

    # Выбираем нужный столбец (O или P)
    idx_col = excel_col_to_index("O" if basis == "start" else "P")
    if not (0 <= idx_col < len(df.columns)):
        return None

    try:
        ser_raw = df.iloc[:, idx_col]
    except Exception:
        return None

    # Приводим значения к датам
    dates = ser_raw.apply(_parse_final_date).dropna()
    if dates.empty:
        return None

    last_date = max(dates)
    if mode == "week":
        start = last_date - timedelta(days=7)
    else:
        # по умолчанию считаем месяц как 30 дней
        start = last_date - timedelta(days=30)
    end = last_date
    return start, end


def build_final_checks_text_filtered(
    df: pd.DataFrame,
    start_date: Optional[date] = None,
    end_date: Optional[date] = None,
    case_no: Optional[str] = None,
    header: str = "📋 Итоговые проверки",
    basis: str = "any",  # "start" / "end" / "any"
) -> str:
    """
    Универсальный вывод итоговых проверок:
    - фильтр по периоду (start_date / end_date) по выбранной базе (O или P);
    - фильтр по номеру дела (case_no).
    """
    df_f = filter_final_checks_df(
        df,
        start_date=start_date,
        end_date=end_date,
        case_no=case_no,
        basis=basis,
    )

    idx_case = excel_col_to_index("B")
    idx_obj = excel_col_to_index("D")
    idx_addr = excel_col_to_index("E")
    idx_start = excel_col_to_index("O")
    idx_end = excel_col_to_index("P")

    lines: List[str] = [header, ""]

    if df_f.empty:
        if case_no:
            return (
                f"По номеру дела {case_no} в таблице итоговых проверок ничего не найдено."
            )
        if start_date and end_date:
            return (
                f"За период {start_date:%d.%m.%Y} — {end_date:%d.%m.%Y} "
                f"итоговые проверки не найдены."
            )
        return "В таблице итоговых проверок нет строк с заполненным номером дела (B)."

    for _, row in df_f.iterrows():

        def safe_text(idx: int) -> str:
            try:
                val = row.iloc[idx]
            except Exception:
                return ""
            if pd.isna(val):
                return ""
            return str(val).strip()

        case_val = safe_text(idx_case)
        if not case_val:
            continue

        obj = safe_text(idx_obj)
        addr = safe_text(idx_addr)

        d_start_raw = row.iloc[idx_start] if idx_start < len(row) else None
        d_end_raw = row.iloc[idx_end] if idx_end < len(row) else None

        row_start = _parse_final_date(d_start_raw)
        row_end = _parse_final_date(d_end_raw)

        def fmt_date(d: Optional[date]) -> str:
            return d.strftime("%d.%m.%Y") if d else ""

        d_start = fmt_date(row_start)
        d_end = fmt_date(row_end)

        lines.append(f"Номер дела: {case_val}")
        if obj:
            lines.append(f"Объект: {obj}")
        if addr:
            lines.append(f"Адрес: {addr}")
        if d_start or d_end:
            if d_start and d_end:
                lines.append(f"Период итоговой проверки: {d_start} — {d_end}")
            elif d_start:
                lines.append(f"Дата начала итоговой проверки: {d_start}")
            else:
                lines.append(f"Дата окончания итоговой проверки: {d_end}")
        lines.append("")
        lines.append("────────────")
        lines.append("")

    return "\n".join(lines)


def build_final_checks_text(df: pd.DataFrame) -> str:
    """
    Старый интерфейс (без фильтров) — на всякий случай.
    """
    return build_final_checks_text_filtered(df)


# -------------------------------------------------
# Итоговые проверки: «BI‑панель» по нарушениям (1–10 дней)
# -------------------------------------------------

FINAL_CHECKS_TARGET_SHEETS_HINTS = [
    ("мкд", ["2025", "мкд"]),
    ("соцобъекты", ["2025", "соц", "объект"]),
    ("остальное", ["2025", "остал"]),
]


def _pick_final_checks_target_sheets(all_sheets: List[str]) -> List[str]:
    """
    Выбирает нужные листы итоговых проверок по «мягким» правилам,
    чтобы не зависеть от точных пробелов в названии.
    Требуемые листы по постановке:
      • 2025 ... МКД
      • 2025 ... СОЦОБЪЕКТЫ
      • 2025 ... Остальное
    """
    if not all_sheets:
        return []

    picked: List[str] = []
    for sheet in all_sheets:
        s = str(sheet).lower().replace("\xa0", " ").strip()
        s_compact = re.sub(r"\s+", " ", s)

        # МКД
        if ("2025" in s_compact) and ("мкд" in s_compact):
            picked.append(sheet)
            continue

        # СОЦОБЪЕКТЫ (разные написания)
        if ("2025" in s_compact) and ("соц" in s_compact) and ("объект" in s_compact):
            picked.append(sheet)
            continue

        # Остальное
        if ("2025" in s_compact) and ("остал" in s_compact):
            picked.append(sheet)
            continue

    # Убираем дубликаты, сохраняем порядок
    uniq: List[str] = []
    for x in picked:
        if x not in uniq:
            uniq.append(x)
    return uniq


def get_final_checks_df_target_sheets() -> Optional[pd.DataFrame]:
    """
    Читает локальный файл итоговых проверок и склеивает только целевые листы:
    «...МКД», «...СОЦОБЪЕКТЫ», «...Остальное» за 2025 год.
    """
    path = FINAL_CHECKS_LOCAL_PATH
    if not path or not os.path.exists(path):
        log.error("Локальный файл итоговых проверок не найден: %s", path)
        return None

    try:
        xls = pd.ExcelFile(path)
        target_sheets = _pick_final_checks_target_sheets(xls.sheet_names)

        if not target_sheets:
            log.warning(
                "Не удалось определить целевые листы итоговых проверок (МКД/СОЦОБЪЕКТЫ/Остальное). "
                "Будут использованы все листы."
            )
            target_sheets = list(xls.sheet_names)

        frames: List[pd.DataFrame] = []
        for sh in target_sheets:
            try:
                df_sh = pd.read_excel(xls, sheet_name=sh)
                df_sh = df_sh.dropna(how="all")
                if not df_sh.empty:
                    df_sh["_final_sheet"] = sh
                    frames.append(df_sh)
            except Exception as e:
                log.warning("Ошибка чтения листа '%s' итоговых проверок: %s", sh, e)

        if not frames:
            return pd.DataFrame()

        df = pd.concat(frames, ignore_index=True).reset_index(drop=True)
        return df
    except Exception as e:
        log.error("Ошибка чтения итоговых проверок (целевые листы): %s", e)
        return None


def _cell_has_net(val: Any) -> bool:
    """
    Проверяет наличие слова «нет» (в любом регистре) в ячейке.
    Учитывает варианты вроде «нет», «нет.» «нет/…», «Нет».
    """
    if val is None:
        return False
    try:
        if isinstance(val, float) and pd.isna(val):
            return False
    except Exception:
        pass
    s = str(val).replace("\n", " ").strip().lower()
    if not s:
        return False
    return "нет" in s


def build_final_checks_violations_bi_panel(
    df_final: pd.DataFrame,
    df_remarks: pd.DataFrame,
    days_min: int = 1,
    days_max: int = 10,
    report_day: Optional[date] = None,
) -> str:
    """
    Формирует «BI‑панель»:
    1) Берёт итоговые проверки из df_final (колонка O — дата начала, колонка B — номер дела).
    2) Отбирает дела, у которых до даты начала (O) осталось 1–10 дней.
    3) Для каждого номера дела ищет строку в таблице замечаний (колонка I).
    4) Проверяет колонки Q, R, Y, AD на наличие слова «нет».
    5) Показывает только дела, где есть хотя бы одно «нет».

    Вывод — по образцу пользователя: номер дела и перечень пунктов «… — нет».
    """
    if report_day is None:
        report_day = local_now().date()

    if df_final is None or df_final.empty:
        return "В таблице итоговых проверок нет данных."

    if df_remarks is None or df_remarks.empty:
        return "Не удалось открыть таблицу замечаний для проверки статусов «нет»."

    idx_case_f = excel_col_to_index("B")
    idx_date_o = excel_col_to_index("O")

    if not (0 <= idx_case_f < len(df_final.columns)) or not (0 <= idx_date_o < len(df_final.columns)):
        return (
            "Не удалось определить обязательные колонки в итоговых проверках "
            "(B — номер дела, O — дата начала)."
        )

    # 1) Отбираем дела на 1–10 дней вперёд по дате O
    candidates: Dict[str, Dict[str, Any]] = {}

    for _, row in df_final.iterrows():
        try:
            case_raw = row.iloc[idx_case_f]
        except Exception:
            continue

        case_norm = normalize_case_number(case_raw)
        if not case_norm:
            continue

        try:
            o_raw = row.iloc[idx_date_o]
        except Exception:
            o_raw = None

        o_date = _parse_final_date(o_raw)
        if not o_date:
            continue

        delta = (o_date - report_day).days
        if delta < days_min or delta > days_max:
            continue

        # сохраняем минимальную дату (если дубли)
        prev = candidates.get(case_norm)
        if not prev or o_date < prev["o_date"]:
            candidates[case_norm] = {
                "case": case_norm,
                "o_date": o_date,
                "delta": delta,
                "sheet": str(row.get("_final_sheet", "")).strip(),
            }

    if not candidates:
        return (
            f"Итоговые проверки на {report_day:%d.%m.%Y}:\n"
            f"В период {days_min}–{days_max} дней (по дате начала O) дел не найдено."
        )

    # 2) Индексы и нормализация для таблицы замечаний (вторая таблица)
    idx_case_r = excel_col_to_index("I")
    if not (0 <= idx_case_r < len(df_remarks.columns)):
        return (
            "Не удалось определить колонку I (Номер дела) во второй таблице (замечания)."
        )

    idx_q = excel_col_to_index("Q")
    idx_r = excel_col_to_index("R")
    idx_y = excel_col_to_index("Y")
    idx_ad = excel_col_to_index("AD")

    # Предварительно считаем нормализованный номер дела для всех строк замечаний
    try:
        remarks_case_norm = df_remarks.iloc[:, idx_case_r].apply(normalize_case_number)
    except Exception:
        remarks_case_norm = pd.Series([""] * len(df_remarks), index=df_remarks.index)

    ISSUE_TITLES = [
        ("Q", "Отметка об устранении замечаний ПБ"),
        ("R", "Отметка об устранении замечаний ПБ в ЗК КНД"),
        ("Y", "Отметка об устранении нарушений АР, ММГН, АГО"),
        ("AD", "Отметка об устранении нарушений ЭОМ"),
    ]

    idx_map = {"Q": idx_q, "R": idx_r, "Y": idx_y, "AD": idx_ad}

    # 3) Собираем только нарушения (есть «нет»)
    out_items: List[Dict[str, Any]] = []

    # сортируем по дате начала O
    ordered = sorted(candidates.values(), key=lambda x: x["o_date"])

    for item in ordered:
        case_norm = item["case"]
        mask = remarks_case_norm == case_norm
        if not mask.any():
            continue  # если в замечаниях нет дела — не считаем его «нарушением»

        df_case = df_remarks.loc[mask]

        issues_present: List[str] = []
        for col_key, title in ISSUE_TITLES:
            idx_col = idx_map.get(col_key)
            if idx_col is None or idx_col >= len(df_remarks.columns):
                continue

            has_net = False
            for _, rrow in df_case.iterrows():
                try:
                    v = rrow.iloc[idx_col]
                except Exception:
                    v = None
                if _cell_has_net(v):
                    has_net = True
                    break
            if has_net:
                issues_present.append(title)

        if not issues_present:
            continue

        item_out = dict(item)
        item_out["issues"] = issues_present
        out_items.append(item_out)

    if not out_items:
        return (
            f"Итоговые проверки на {report_day:%d.%m.%Y}:\n"
            f"В период {days_min}–{days_max} дней (по дате начала O) "
            f"дела есть, но по ним не найдено статусов «нет» в колонках Q/R/Y/AD."
        )

    # 4) Финальный текст
    lines: List[str] = []
    lines.append(f"Итоговые проверки на {report_day:%d.%m.%Y}:")
    lines.append(f"Период: {days_min}–{days_max} дней (по дате начала O).")
    lines.append("")

    for it in out_items:
        case_no = it["case"]
        lines.append(f"{case_no}")

        issues = it.get("issues") or []
        for n, title in enumerate(issues, start=1):
            lines.append(f"{n}) {title} - нет")
        lines.append("")

    return "\n".join(lines).strip()






def build_final_checks_kpi_dashboard(
    df_final: pd.DataFrame,
    df_remarks: Optional[pd.DataFrame],
    days_min: int = 1,
    days_max: int = 10,
) -> str:
    """
    Строит короткий дашборд по итоговым проверкам на ближайший интервал дней.

    Логика отбора:
    - берётся столбец O «Дата начала итоговой проверки (в формате дд/мм/гггг)»;
    - дата парсится через pandas.to_datetime c dayfirst=True;
    - считаем разницу (в днях) между датой начала итоговой проверки и сегодняшним днём;
    - в выборку попадают только дела, у которых days_min <= delta_days <= days_max.

    Параллельно, по листу с замечаниями (ПБ, АР, ММГН, АГО) определяется,
    есть ли по делу хотя бы одно значение «нет» в ключевых колонках (Q, R, Y, AE).
    """
    today = date.today()
    lines: List[str] = []

    lines.append("📋 Раздел «Итоговые проверки»")
    lines.append("")

    # 1. Базовые проверки
    if df_final is None or df_final.empty:
        lines.append("В таблице итоговых проверок нет данных.")
        return "\n".join(lines)

    # Определяем нужные столбцы в итоговой таблице
    idx_case = get_col_index_by_header(df_final, "номер дела", "B")
    idx_address = get_col_index_by_header(df_final, "адрес объекта", "F")
    idx_start = get_col_index_by_header(df_final, "дата начала итоговой проверки", "O")
    idx_end = get_col_index_by_header(df_final, "дата окончания итоговой проверки", "P")

    if idx_case is None or idx_start is None:
        lines.append(
            "Не удалось найти обязательные колонки (номер дела или дата начала итоговой проверки)."
        )
        return "\n".join(lines)

    col_case = df_final.columns[idx_case]
    col_address = df_final.columns[idx_address] if idx_address is not None else None
    col_start = df_final.columns[idx_start]
    col_end = df_final.columns[idx_end] if idx_end is not None else None

    # 2. Нормализуем даты начала итоговой проверки через pandas
    start_series = pd.to_datetime(df_final[col_start], errors="coerce", dayfirst=True)
    if start_series.notna().sum() == 0:
        lines.append(
            "Не удалось корректно распознать даты начала итоговой проверки в столбце O."
        )
        return "\n".join(lines)

    today_ts = pd.to_datetime(today)
    deltas = (start_series - today_ts).dt.days

    mask = (deltas >= days_min) & (deltas <= days_max)
    df_window = df_final.loc[mask].copy()

    if df_window.empty:
        lines.append(
            f"В ближайшие {days_min}–{days_max} дней (по дате начала итоговой проверки, столбец O) "
            f"нет дел с запланированной итоговой проверкой."
        )
        return "\n".join(lines)

    # Для сортировки по дате
    df_window["_start_ts"] = pd.to_datetime(
        df_window[col_start], errors="coerce", dayfirst=True
    )
    df_window.sort_values("_start_ts", inplace=True)

    # 3. Собираем по листу замечаний номера дел, где есть хотя бы одно «нет»
    cases_with_not_fixed: set[str] = set()
    if df_remarks is not None and not df_remarks.empty:
        idx_case_r = get_col_index_by_header(df_remarks, "номер дела", "I")
        idx_pb_r = get_col_index_by_header(df_remarks, "пожар", "Q")
        idx_ar_r = get_col_index_by_header(df_remarks, "архит", "R")
        idx_mgn_r = get_col_index_by_header(df_remarks, "мгн", "Y")
        idx_ago_r = get_col_index_by_header(df_remarks, "аго", "AE")

        def _is_net(val: Any) -> bool:
            if val is None:
                return False
            if isinstance(val, float) and pd.isna(val):
                return False
            s = str(val).strip().lower()
            return s == "нет"

        if idx_case_r is not None:
            for _, row in df_remarks.iterrows():
                case_num = str(row.iloc[idx_case_r]).strip()
                if not case_num:
                    continue

                has_net = False
                for idx_col in (idx_pb_r, idx_ar_r, idx_mgn_r, idx_ago_r):
                    if idx_col is not None and _is_net(row.iloc[idx_col]):
                        has_net = True
                        break

                if has_net:
                    cases_with_not_fixed.add(case_num)

    # 4. Статистика по окну дат
    total_cases = len(df_window)

    # Считаем, сколько из этих дел имеют неустранённые замечания
    cases_in_window = set()
    cases_in_window_with_not_fixed = set()

    for _, row in df_window.iterrows():
        case_num = str(row[col_case]).strip()
        if not case_num:
            continue
        cases_in_window.add(case_num)
        if case_num in cases_with_not_fixed:
            cases_in_window_with_not_fixed.add(case_num)

    lines.append(
        f"В ближайшие {days_min}–{days_max} дней (по дате начала итоговой проверки, столбец O) "
        f"запланировано дел: {total_cases}."
    )

    if cases_in_window_with_not_fixed:
        lines.append(
            "Из них с неустранёнными замечаниями по листу «ПБ, АР,ММГН, АГО (2025)»: "
            f"{len(cases_in_window_with_not_fixed)}."
        )
    else:
        lines.append("По данным листа с замечаниями в этом окне нет дел с неустранёнными замечаниями.")

    # 5. Перечень дел
    lines.append("")
    lines.append("Перечень дел (максимум 40 строк):")

    max_list = 40
    printed = 0

    for _, row in df_window.iterrows():
        case_num = str(row[col_case]).strip()
        if not case_num:
            continue

        addr = str(row[col_address]).strip() if col_address is not None else ""
        start_ts = row["_start_ts"]
        start_str = start_ts.strftime("%d.%m.%Y") if not pd.isna(start_ts) else "?"

        mark = "❌" if case_num in cases_with_not_fixed else "✅"
        line = f"{mark} {case_num} — начало итоговой проверки {start_str}"
        if addr:
            line += f"; {addr}"

        lines.append(line)
        printed += 1
        if printed >= max_list:
            if total_cases > printed:
                lines.append(f"... и ещё {total_cases - printed} дел.")
            break

    return "\n".join(lines)


async def send_final_checks_xlsx_filtered(
    chat_id: int,
    df: pd.DataFrame,
    context: ContextTypes.DEFAULT_TYPE,
    start_date: Optional[date] = None,
    end_date: Optional[date] = None,
    case_no: Optional[str] = None,
    filename_suffix: str = "",
    basis: str = "any",
):
    df_f = filter_final_checks_df(
        df,
        start_date=start_date,
        end_date=end_date,
        case_no=case_no,
        basis=basis,
    )
    if df_f.empty:
        await context.bot.send_message(
            chat_id=chat_id,
            text="Нет данных для выгрузки по выбранным условиям.",
        )
        return

    bio = BytesIO()
    df_f.to_excel(bio, sheet_name="Итоговые проверки", index=False)
    bio.seek(0)

    fname = "Итоговые_проверки"
    parts = []
    if case_no:
        parts.append(f"дело_{case_no}")
    if start_date and end_date:
        parts.append(f"{start_date:%d.%m.%Y}-{end_date:%d.%m.%Y}")
    if filename_suffix:
        parts.append(filename_suffix)
    if parts:
        fname += "_" + "_".join(parts)
    fname += ".xlsx"

    await context.bot.send_document(
        chat_id=chat_id,
        document=InputFile(bio, filename=fname),
        caption="Итоговые проверки (фильтрованный список)",
    )


# -------------------------------------------------
# Инспектор → Google Sheets
# -------------------------------------------------
def append_inspector_row_to_excel(form: Dict[str, Any]) -> bool:
    service = get_sheets_service()
    if service is None:
        log.error("Google Sheets API недоступен.")
        return False

    try:
        area_str = str(form.get("area", "")).replace(".", ",")
        floors_str = str(form.get("floors", ""))

        d_value = (
            f"Площадь (кв.м): {area_str}\n"
            f"Количество этажей: {floors_str}"
        )

        row = [
            "",
            form.get("date").strftime("%d.%m.%Y") if form.get("date") else "",
            "",
            d_value,
            form.get("onzs", ""),
            form.get("developer", ""),
            form.get("object", ""),
            form.get("address", ""),
            form.get("case", ""),
            form.get("check_type", ""),
        ]

        body = {"values": [row]}

        response = (
            service.spreadsheets()
            .values()
            .append(
                spreadsheetId=GSHEETS_SPREADSHEET_ID,
                range=f"'{INSPECTOR_SHEET_NAME}'!A1",
                valueInputOption="USER_ENTERED",
                insertDataOption="INSERT_ROWS",
                body=body,
            )
            .execute()
        )

        log.info("Инспектор: запись добавлена в Google Sheets: %s", response)
        return True

    except Exception as e:
        log.error("Ошибка записи инспектора в Google Sheets: %s", e)
        return False


# -------------------------------------------------
# Инспектор — мастер
# -------------------------------------------------
async def inspector_process(update: Update, context: ContextTypes.DEFAULT_TYPE):
    text = update.message.text
    form = context.user_data.get("inspector_form", {}) or {}
    step = form.get("step")

    if not step:
        context.user_data["inspector_form"] = {"step": "date"}
        await update.message.reply_text(
            "👮‍♂️ Выезд инспектора\n\n"
            "1/8. Дата выезда (ДД.ММ.ГГГГ):"
        )
        return

    if step == "date":
        try:
            form["date"] = datetime.strptime(text, "%d.%m.%Y").date()
            form["step"] = "area"
            context.user_data["inspector_form"] = form
            await update.message.reply_text("1/8. Площадь объекта (кв.м):")
        except Exception:
            await update.message.reply_text(
                "Введите дату в формате ДД.ММ.ГГГГ (например, 30.12.2025)"
            )
        return

    if step == "area":
        form["area"] = text
        form["step"] = "floors"
        context.user_data["inspector_form"] = form
        await update.message.reply_text("2/8. Количество этажей:")
        return

    if step == "floors":
        form["floors"] = text
        form["step"] = "onzs"
        context.user_data["inspector_form"] = form
        await update.message.reply_text("3/8. ОНзС (1–12):")
        return

    if step == "onzs":
        form["onzs"] = text
        form["step"] = "developer"
        context.user_data["inspector_form"] = form
        await update.message.reply_text("4/8. Наименование застройщика:")
        return

    if step == "developer":
        form["developer"] = text
        form["step"] = "object"
        context.user_data["inspector_form"] = form
        await update.message.reply_text("5/8. Наименование объекта:")
        return

    if step == "object":
        form["object"] = text
        form["step"] = "address"
        context.user_data["inspector_form"] = form
        await update.message.reply_text("6/8. Строительный адрес:")
        return

    if step == "address":
        form["address"] = text
        form["step"] = "case"
        context.user_data["inspector_form"] = form
        await update.message.reply_text("7/8. Номер дела (формат 00-00-000000):")
        return

    if step == "case":
        form["case"] = text
        form["step"] = "check_type"
        context.user_data["inspector_form"] = form
        await update.message.reply_text(
            "8/8. Вид проверки (ПП, итоговая, профвизит, поручение и т.п.):"
        )
        return

    if step == "check_type":
        form["check_type"] = text
        form["step"] = "done"
        context.user_data["inspector_form"] = form

        await update.message.reply_text("⏳ Сохраняю выезд...")

        ok_db = save_inspector_to_db(form)
        ok_gs = append_inspector_row_to_excel(form)

        if ok_db and ok_gs:
            msg = "✅ Выезд сохранён в боте и добавлен в общую таблицу."
        elif ok_db and not ok_gs:
            msg = (
                "✅ Выезд сохранён в боте.\n"
                "⚠ Не удалось добавить в Google Sheets (проверьте ключ/права)."
            )
        elif not ok_db and ok_gs:
            msg = (
                "⚠ Выезд добавлен в Google Sheets, но не удалось сохранить локную запись."
            )
        else:
            msg = (
                "❌ Не удалось сохранить выезд ни локно, ни в Google Sheets.\n"
                "Сообщите разработчику."
            )

        await update.message.reply_text(msg)
        context.user_data.pop("inspector_form", None)
        return


# -------------------------------------------------
# ОНзС
# -------------------------------------------------
def onzs_menu_inline() -> InlineKeyboardMarkup:
    buttons = []
    row = []
    for i in range(1, 13):
        row.append(InlineKeyboardButton(str(i), callback_data=f"onzs_filter_{i}"))
        if len(row) == 4:
            buttons.append(row)
            row = []
    if row:
        buttons.append(row)
    return InlineKeyboardMarkup(buttons)


def build_onzs_list_by_number(df: pd.DataFrame, number: str) -> str:
    onzs_idx = get_col_index_by_header(df, "онзс", "D")
    if onzs_idx is None:
        return "Не удалось определить столбец ОНзС в файле замечаний."

    case_idx = get_case_col_index(df)
    addr_idx = get_col_index_by_header(df, "строительный адрес", "H")

    num_str = normalize_onzs_value(number)
    mask: List[bool] = []
    for _, row in df.iterrows():
        try:
            val_raw = row.iloc[onzs_idx]
        except Exception:
            val_raw = None
        val_norm = normalize_onzs_value(val_raw)
        mask.append(val_norm == num_str)

    if not any(mask):
        return f"Нет объектов с ОНзС = {number}."

    df_f = df[mask]

    lines = [f"ОНзС = {number}", f"Найдено дел: {len(df_f)}", ""]

    for _, row in df_f.iterrows():

        def safe(idx: Optional[int]) -> str:
            if idx is None:
                return ""
            try:
                val = row.iloc[idx]
            except Exception:
                return ""
            try:
                if pd.isna(val):
                    return ""
            except Exception:
                pass
            s = str(val).strip()
            if not s or s.lower() == "nan":
                return ""
            return s

        case_no = safe(case_idx)
        addr = safe(addr_idx)

        if not case_no and not addr:
            continue

        if case_no and addr:
            lines.append(f"• {case_no} — {addr}")
        elif case_no:
            lines.append(f"• {case_no}")
        else:
            lines.append(f"• {addr}")

    return "\n".join(lines)


# -------------------------------------------------
# Инспектор — список/Excel
# -------------------------------------------------
def build_inspector_list_text(rows: List[sqlite3.Row]) -> str:
    if not rows:
        return "Пока нет сохранённых выездов инспектора."

    lines: List[str] = ["Последние выезды инспектора:", ""]
    for r in rows:
        d = r["date"] or ""
        try:
            d_fmt = datetime.strptime(d, "%Y-%m-%d").strftime("%d.%m.%Y")
        except Exception:
            d_fmt = d
        lines.append(
            f"• {d_fmt} — дело {r['case_no'] or '-'}, "
            f"ОНзС {r['onzs'] or '-'}, {r['check_type'] or ''}"
        )
        addr = r["address"] or ""
        if addr:
            lines.append(f"  Адрес: {addr}")
        obj = r["object"] or ""
        if obj:
            lines.append(f"  Объект: {obj}")
        dev = r["developer"] or ""
        if dev:
            lines.append(f"  Застройщик: {dev}")
        lines.append("")
    return "\n".join(lines)


async def send_inspector_xlsx(
    chat_id: int, rows: List[sqlite3.Row], context: ContextTypes.DEFAULT_TYPE
):
    if not rows:
        await context.bot.send_message(
            chat_id=chat_id, text="Пока нет сохранённых выездов инспектора."
        )
        return

    data = []
    for r in rows:
        d = r["date"] or ""
        try:
            d_fmt = datetime.strptime(d, "%Y-%m-%d").strftime("%d.%m.%Y")
        except Exception:
            d_fmt = d
        data.append(
            {
                "Дата выезда": d_fmt,
                "Площадь (кв.м)": r["area"] or "",
                "Этажность": r["floors"] or "",
                "ОНзС": r["onzs"] or "",
                "Застройщик": r["developer"] or "",
                "Наименование объекта": r["object"] or "",
                "Строительный адрес": r["address"] or "",
                "Номер дела": r["case_no"] or "",
                "Вид проверки": r["check_type"] or "",
            }
        )

    df = pd.DataFrame(data)

    bio = BytesIO()
    with pd.ExcelWriter(bio, engine="openpyxl") as writer:
        df.to_excel(writer, sheet_name="Инспектор", index=False)

    bio.seek(0)
    filename = f"Инспектор_выезды_{date.today().strftime('%d.%m.%Y')}.xlsx"

    await context.bot.send_document(
        chat_id=chat_id,
        document=InputFile(bio, filename=filename),
        caption="Выезды инспектора (отдельный файл)",
    )


# -------------------------------------------------
# CALLBACK HANDLER
# -------------------------------------------------
async def callback_handler(update: Update, context: ContextTypes.DEFAULT_TYPE):
    query = update.callback_query
    data = query.data
    user = query.from_user
    await query.answer()

    # ТЗ для ЦНИЛ — inline-кнопки мастера
    if data and str(data).startswith("cnil:"):
        await cnil_callback(update, context)
        return


    settings = get_schedule_state()
    version = get_schedule_version(settings)


    # --- 🚨 КРАСНЫЕ ЛАМПОЧКИ ---
    if data == "redlamps_upload":
        context.user_data["awaiting_redlamps_upload"] = True
        await query.message.reply_text(
            "📤 Отправьте Excel-файл (.xlsx) для раздела «Красные лампочки»."
        )
        return

    if data == "redlamps_period":
        context.user_data["redlamps_period"] = {"step": "start"}
        await query.message.reply_text("Введите дату начала периода (ДД.ММ.ГГГГ):")
        return

    if data == "redlamps_reset":
        context.user_data.pop("redlamps_file_bytes", None)
        context.user_data.pop("redlamps_file_name", None)
        context.user_data.pop("redlamps_period", None)
        context.user_data.pop("awaiting_redlamps_upload", None)
        await query.message.reply_text("Сброс выполнен. Загрузите файл заново при необходимости.")
        return

    if data == "redlamps_build":
        xbytes = context.user_data.get("redlamps_file_bytes")
        if not xbytes:
            await query.message.reply_text("Сначала загрузите Excel-файл кнопкой «📤 Загрузить файл».")
            return

        period = context.user_data.get("redlamps_period") or {}
        d_from = period.get("date_from")
        d_to = period.get("date_to")
        if not d_from or not d_to:
            await query.message.reply_text("Сначала задайте период кнопкой «📅 Выбрать период (K–L)».")
            return

        try:
            text_out = _redlamps_process_bytes(
                xlsx_bytes=xbytes,
                date_from=d_from,
                date_to=d_to,
                tolerance_days=REDLAMPS_TOLERANCE_DAYS_DEFAULT,
            )
        except Exception as e:
            log.error("REDLAMPS: ошибка обработки файла: %s", e)
            await query.message.reply_text("Ошибка обработки файла. Проверьте структуру Excel.")
            return

        await send_long_text(query.message.chat, text_out)
        return
    # --- ГРАФИК ---
    if data == "schedule_refresh":
        df = get_schedule_df()
        if df is None:
            await query.message.reply_text("Не удалось прочитать лист «График».")
        else:
            await query.message.reply_text(f"Лист «График» прочитан, строк: {len(df)}.")
        return

    if data == "schedule_download":
        df = get_schedule_df()
        if df is None or df.empty:
            await query.message.reply_text(
                "Не удалось получить лист «График» для выгрузки."
            )
            return

        await send_schedule_xlsx(
            chat_id=query.message.chat.id,
            dataframe=df,
            context=context,
        )
        return

    if data == "schedule_upload":
        await query.message.reply_text("Загрузка графика в этой сборке не реализована.")
        return

    if data == "schedule_approvers":
        if not is_admin(user.id):
            await query.message.reply_text(
                "Только администратор может настраивать согласующих."
            )
            return
        context.user_data["awaiting_approvers_input"] = {"version": version}
        await query.message.reply_text(
            "Отправьте список согласующих (юзернеймы через пробел/запятую/новую строку), например:\n"
            "@asdinamitif @FrolovAlNGSN @cappit_G59"
        )
        return

    if data.startswith("schedule_approve:") or data.startswith("schedule_rework:"):
        action, approver_tag = data.split(":", 1)
        user_username = user.username or ""
        user_tag = f"@{user_username}" if user_username else ""

        if user_tag.lower() != approver_tag.lower():
            await query.answer(
                text=f"Эта кнопка предназначена для {approver_tag}.",
                show_alert=True,
            )
            return

        if action == "schedule_approve":
            update_schedule_approval_status(version, approver_tag, "approved", None)
            await query.message.reply_text(
                f"{approver_tag} согласовал(а) график. Спасибо!"
            )

            approvals = get_schedule_approvals(version)
            if approvals and all(r["status"] == "approved" for r in approvals):
                header = build_schedule_header(version, approvals)
                lines = [header, "", "Согласовано всеми:"]
                for r in approvals:
                    lines.append(
                        f"• {r['approver']} — {_format_dt(r['decided_at'])} ✅"
                    )
                text = "\n".join(lines)

                write_schedule_summary_to_sheet(version, approvals)

                if SCHEDULE_NOTIFY_CHAT_ID is not None:
                    try:
                        await context.bot.send_message(
                            chat_id=SCHEDULE_NOTIFY_CHAT_ID, text=text
                        )
                    except Exception as e:
                        log.error(
                            "Ошибка отправки графика в канал %s: %s",
                            SCHEDULE_NOTIFY_CHAT_ID,
                            e,
                        )
            return

        if action == "schedule_rework":
            context.user_data["awaiting_rework_comment"] = {
                "version": version,
                "approver": approver_tag,
            }
            await query.message.reply_text(
                "Напишите комментарий, почему график нужно доработать."
            )
            return

    # --- ЗАМЕЧАНИЯ ---
    if data == "remarks_search_case":
        context.user_data["awaiting_case_search"] = True
        await query.message.reply_text(
            "Введите номер дела (формат 00-00-000000), который нужно найти:"
        )
        return

    if data == "remarks_onzs":
        kb = onzs_menu_inline()
        msg = (
            "🏗 Раздел «ОНзС»\n\n"
            "Выберите номер ОНзС, чтобы увидеть список дел (Номер дела (I) + адрес) "
            "из текущего файла замечаний.\n"
            "Для выбранного ОНзС можно отдельно показать только неустранённые замечания."
        )
        await query.message.reply_text(msg, reply_markup=kb)
        return

    if data == "remarks_not_done":
        await query.message.reply_text("Ищу строки со статусом «нет»...")
        df = get_remarks_df_current()
        if df is None:
            await query.message.reply_text(
                "Не удалось получить файл замечаний. Проверьте доступ к таблице."
            )
            return
        text = build_remarks_not_done_text(df)
        await send_long_text(query.message.chat, text)
        return

    if data == "remarks_download":
        await query.message.reply_text(
            "Файл с замечаниями и графиком можно открыть по ссылке:\n"
            f"{GOOGLE_SHEET_URL_DEFAULT}"
        )
        return

    if data.startswith("onzs_filter_"):
        number = data.replace("onzs_filter_", "")
        df = get_remarks_df_current()
        if df is None:
            await query.message.reply_text("Не удалось открыть таблицу ОНзС.")
            return
        text = build_onzs_list_by_number(df, number)
        await send_long_text(query.message.chat, text)

        kb = InlineKeyboardMarkup(
            [
                [
                    InlineKeyboardButton(
                        f"❌ Не устранены (ОНзС {number})",
                        callback_data=f"onzs_not_done_{number}",
                    )
                ]
            ]
        )
        await query.message.reply_text(
            f"Для ОНзС {number} можно показать только строки, где статус «нет».",
            reply_markup=kb,
        )
        return

    if data.startswith("onzs_not_done_"):
        number = data.replace("onzs_not_done_", "")
        df = get_remarks_df_current()
        if df is None:
            await query.message.reply_text(
                "Не удалось получить файл замечаний. Проверьте доступ к таблице."
            )
            return
        text = build_remarks_not_done_by_onzs(df, number)
        await send_long_text(query.message.chat, text)
        return

    # --- ИНСПЕКТОР ---
    if data == "inspector_add":
        context.user_data["inspector_form"] = {"step": "date"}
        await query.message.reply_text(
            "👮‍♂️ Выезд инспектора\n\n"
            "Укажем данные по шагам.\n"
            "1/8. Дата выезда (ДД.ММ.ГГГГ):"
        )
        return

    if data == "inspector_list":
        rows = fetch_inspector_visits(limit=50)
        text = build_inspector_list_text(rows)
        await send_long_text(query.message.chat, "\n".join(text.split("\n")))
        return

    if data == "inspector_download":
        rows = fetch_inspector_visits(limit=1000)
        await send_inspector_xlsx(
            chat_id=query.message.chat.id, rows=rows, context=context
        )
        return

    if data == "inspector_reset":
        clear_inspector_visits()
        await query.message.reply_text(
            "Список выездов инспектора очищен.\n"
            "Новые выезды будут попадать в Excel после добавления через кнопку «➕ Добавить выезд»."
        )
        return

    # --- ИТОГОВЫЕ ПРОВЕРКИ ---
    if data == "final_week":
        # запоминаем режим и спрашиваем, по какой дате фильтровать
        context.user_data["final_range_choice"] = {"mode": "week"}
        kb = InlineKeyboardMarkup(
            [
                [
                    InlineKeyboardButton(
                        "📌 По дате начала (O)", callback_data="final_basis_start"
                    ),
                    InlineKeyboardButton(
                        "📌 По дате окончания (P)", callback_data="final_basis_end"
                    ),
                ]
            ]
        )
        await query.message.reply_text(
            "За неделю: по какой дате фильтровать?\n\n"
            "• O — дата начала итоговой проверки\n"
            "• P — дата окончания итоговой проверки",
            reply_markup=kb,
        )
        return

    if data == "final_month":
        context.user_data["final_range_choice"] = {"mode": "month"}
        kb = InlineKeyboardMarkup(
            [
                [
                    InlineKeyboardButton(
                        "📌 По дате начала (O)", callback_data="final_basis_start"
                    ),
                    InlineKeyboardButton(
                        "📌 По дате окончания (P)", callback_data="final_basis_end"
                    ),
                ]
            ]
        )
        await query.message.reply_text(
            "За месяц: по какой дате фильтровать?\n\n"
            "• O — дата начала итоговой проверки\n"
            "• P — дата окончания итоговой проверки",
            reply_markup=kb,
        )
        return

    if data == "final_period":
        context.user_data["final_range_choice"] = {"mode": "period"}
        kb = InlineKeyboardMarkup(
            [
                [
                    InlineKeyboardButton(
                        "📌 По дате начала (O)", callback_data="final_basis_start"
                    ),
                    InlineKeyboardButton(
                        "📌 По дате окончания (P)", callback_data="final_basis_end"
                    ),
                ]
            ]
        )
        await query.message.reply_text(
            "Выбор периода: по какой дате фильтровать?\n\n"
            "• O — дата начала итоговой проверки\n"
            "• P — дата окончания итоговой проверки",
            reply_markup=kb,
        )
        return

    # выбор базы: O или P
    if data in ("final_basis_start", "final_basis_end"):
        basis = "start" if data == "final_basis_start" else "end"
        state = context.user_data.get("final_range_choice")
        if not state:
            await query.message.reply_text(
                "Сначала выберите режим (за неделю/за месяц/выбрать период) в разделе «Итоговые проверки»."
            )
            return

        mode = state.get("mode")
        # недельный и месячный режимы
        if mode in ("week", "month"):
            df = get_final_checks_df()
            if df is None:
                await query.message.reply_text(
                    "Не удалось открыть таблицу итоговых проверок."
                )
                context.user_data.pop("final_range_choice", None)
                return

            period = compute_auto_period_for_final(df, basis=basis, mode=mode)
            if not period:
                await query.message.reply_text(
                    "В таблице итоговых проверок нет корректных дат в выбранном столбце (O или P)."
                )
                context.user_data.pop("final_range_choice", None)
                return

            start, end = period
            if mode == "week":
                mode_text = "за неделю"
            else:
                mode_text = "за месяц"

            basis_text = (
                "по дате начала (O)" if basis == "start" else "по дате окончания (P)"
            )

            header = (
                f"📋 Итоговые проверки {mode_text} {basis_text}\n"
                f"{start:%d.%m.%Y} — {end:%d.%m.%Y}"
            )
            text_out = build_final_checks_text_filtered(
                df,
                start_date=start,
                end_date=end,
                header=header,
                basis=basis,
            )
            await send_long_text(query.message.chat, text_out)
            await send_final_checks_xlsx_filtered(
                chat_id=query.message.chat.id,
                df=df,
                context=context,
                start_date=start,
                end_date=end,
                basis=basis,
            )
            context.user_data.pop("final_range_choice", None)
            return

        # пользовательский период
        if mode == "period":
            context.user_data["final_period"] = {
                "step": "start",
                "basis": basis,
            }
            context.user_data.pop("final_range_choice", None)
            await query.message.reply_text(
                "Введите дату начала периода (ДД.ММ.ГГГГ):"
            )
            return

        # на всякий случай
        context.user_data.pop("final_range_choice", None)
        await query.message.reply_text(
            "Что-то пошло не так. Попробуйте ещё раз выбрать режим."
        )
        return

    if data == "final_search_case":
        context.user_data["awaiting_final_case_search"] = True
        await query.message.reply_text(
            "Введите номер дела (формат 00-00-000000), который нужно найти "
            "в итоговых проверках:"
        )
        return


# -------------------------------------------------
# TEXT ROUTER
# -------------------------------------------------
async def text_router(update: Update, context: ContextTypes.DEFAULT_TYPE):
    text = update.message.text.strip()
    chat = update.message.chat


    # Ассистент (режим консультации по таблице) — включение
    if ENABLE_ASSISTANT and text == "🗣 Ассистент":
        context.user_data["assistant_mode"] = True
        await update.message.reply_text(
            "🗣 Режим ассистента включён.\n"
            "Можно писать или отправлять голосом.\n\n"
            "Примеры:\n"
            "• Найди 03-46-108600, устранено ли по пожарной\n"
            "• Что по делу 09-27-001100?\n"
            "• Найди по застройщику ИНВЕСТЦЕНТР\n\n"
            "Чтобы выйти: напишите «Выход».",
            reply_markup=main_menu(),
        )
        return

    # Ассистент — обработка всех сообщений в режиме
    if ENABLE_ASSISTANT and context.user_data.get("assistant_mode"):
        await assistant_answer(update.message.chat, context, text, recognized_from_voice=False)
        return

    # ТЗ для ЦНИЛ — запуск мастера из главного меню
    if text in ("🧪 ТЗ для ЦНИЛ", "ТЗ для ЦНИЛ"):
        await cnil_start(update, context)
        return

    # ТЗ для ЦНИЛ — действия раздела (кнопки меню)
    if text == "📝 Заполнить форму":
        await cnil_start_form(update, context)
        return

    if text == "⬇️ Скачать таблицу":
        context.user_data["cnil_download_wait"] = True
        context.user_data.pop("cnil_change_step", None)
        await update.message.reply_text(
            "Введите пароль для скачивания таблицы ЦНИЛ:",
            reply_markup=cnil_menu_keyboard(),
        )
        return

    if text == "🔐 Изменить пароль скачивания":
        context.user_data["cnil_change_step"] = 1
        context.user_data.pop("cnil_download_wait", None)
        await update.message.reply_text(
            "Введите текущий пароль скачивания (или резервный 051995):",
            reply_markup=cnil_menu_keyboard(),
        )
        return

    if text == "⬅️ Назад":
        # очистить состояния раздела
        context.user_data.pop("cnil", None)
        context.user_data.pop("cnil_download_wait", None)
        context.user_data.pop("cnil_change_step", None)
        context.user_data.pop("cnil_new_password", None)
        await start(update, context)
        return

    # ТЗ для ЦНИЛ — ожидание ввода пароля для скачивания
    if context.user_data.get("cnil_download_wait"):
        entered = text.strip()
        ok = (entered == cnil_load_download_password()) or (entered == CNIL_MASTER_DOWNLOAD_PASSWORD)
        if not ok:
            await update.message.reply_text(
                "❌ Неверный пароль. Попробуйте ещё раз или нажмите «⬅️ Назад».",
                reply_markup=cnil_menu_keyboard(),
            )
            return
        context.user_data.pop("cnil_download_wait", None)
        await cnil_send_results_excel(update, context)
        return

    # ТЗ для ЦНИЛ — смена пароля (шаг 1: проверка старого)
    if context.user_data.get("cnil_change_step") == 1:
        entered = text.strip()
        ok = (entered == cnil_load_download_password()) or (entered == CNIL_MASTER_DOWNLOAD_PASSWORD)
        if not ok:
            await update.message.reply_text(
                "❌ Неверный пароль. Введите текущий пароль скачивания (или резервный 051995):",
                reply_markup=cnil_menu_keyboard(),
            )
            return
        context.user_data["cnil_change_step"] = 2
        await update.message.reply_text(
            "Введите новый пароль для скачивания (4–32 символа):",
            reply_markup=cnil_menu_keyboard(),
        )
        return

    # ТЗ для ЦНИЛ — смена пароля (шаг 2: новый пароль)
    if context.user_data.get("cnil_change_step") == 2:
        new_pw = text.strip()
        if len(new_pw) < 4 or len(new_pw) > 32:
            await update.message.reply_text(
                "❌ Пароль должен быть длиной 4–32 символа. Введите новый пароль:",
                reply_markup=cnil_menu_keyboard(),
            )
            return
        if new_pw == CNIL_MASTER_DOWNLOAD_PASSWORD:
            await update.message.reply_text(
                "❌ Этот пароль зарезервирован и не может быть установлен как основной. Введите другой пароль:",
                reply_markup=cnil_menu_keyboard(),
            )
            return
        cnil_save_download_password(new_pw)
        context.user_data.pop("cnil_change_step", None)
        await update.message.reply_text(
            "✅ Пароль скачивания обновлён.\n\nТеперь таблицу можно скачать по новому паролю (или по резервному 051995).",
            reply_markup=cnil_menu_keyboard(),
        )
        return

    # ТЗ для ЦНИЛ — если мастер активен, обрабатываем шаги ввода
    if context.user_data.get("cnil"):
        await cnil_text_step(update, context)
        return

    # Инспектор — пошаговый мастер
    if "inspector_form" in context.user_data:
        await inspector_process(update, context)
        return

    # Итоговые проверки — пользовательский период
    if context.user_data.get("final_period"):
        period = context.user_data["final_period"]
        step = period.get("step")
        basis = period.get("basis", "any")

        # ШАГ 1: ввод даты начала
        if step == "start":
            try:
                start_date = datetime.strptime(text, "%d.%m.%Y").date()
                if start_date.year < 2000 or start_date.year > 2100:
                    raise ValueError("year out of range")

                period["start_date"] = start_date
                period["step"] = "end"
                context.user_data["final_period"] = period
                await update.message.reply_text(
                    "Введите дату окончания периода (ДД.ММ.ГГГГ):"
                )
            except Exception:
                await update.message.reply_text(
                    "Дата начала в неверном формате.\n"
                    "Введите в виде ДД.ММ.ГГГГ (например, 05.01.2025)."
                )
            return

        # ШАГ 2: ввод даты окончания
        if step == "end":
            try:
                end_date = datetime.strptime(text, "%d.%m.%Y").date()
                if end_date.year < 2000 or end_date.year > 2100:
                    raise ValueError("year out of range")

                start_date = period.get("start_date")
                if start_date and end_date < start_date:
                    await update.message.reply_text(
                        "Дата окончания раньше даты начала.\n"
                        "Введите корректную дату окончания (ДД.ММ.ГГГГ)."
                    )
                    return

                df = get_final_checks_df()
                if df is None:
                    await update.message.reply_text(
                        "Не удалось открыть таблицу итоговых проверок."
                    )
                    context.user_data.pop("final_period", None)
                    return

                basis_text = (
                    "по дате начала (O)" if basis == "start" else "по дате окончания (P)"
                )
                header = (
                    f"📋 Итоговые проверки {basis_text} "
                    f"за период {start_date:%d.%m.%Y} — {end_date:%d.%m.%Y}"
                )
                text_out = build_final_checks_text_filtered(
                    df,
                    start_date=start_date,
                    end_date=end_date,
                    header=header,
                    basis=basis,
                )
                await send_long_text(chat, text_out)
                await send_final_checks_xlsx_filtered(
                    chat_id=chat.id,
                    df=df,
                    context=context,
                    start_date=start_date,
                    end_date=end_date,
                    basis=basis,
                )
                context.user_data.pop("final_period", None)
            except Exception:
                await update.message.reply_text(
                    "Дата окончания в неверном формате.\n"
                    "Введите в виде ДД.ММ.ГГГГ (например, 12.12.2025)."
                )
            return

    
    # 🚨 Красные лампочки — ввод периода
    if context.user_data.get("redlamps_period"):
        rp = context.user_data["redlamps_period"]
        step = rp.get("step")

        if step == "start":
            try:
                d_from = datetime.strptime(text, "%d.%m.%Y").date()
                rp["date_from"] = d_from
                rp["step"] = "end"
                context.user_data["redlamps_period"] = rp
                await update.message.reply_text("Введите дату окончания периода (ДД.ММ.ГГГГ):")
            except Exception:
                await update.message.reply_text(
                    "Дата начала в неверном формате. Введите ДД.ММ.ГГГГ (например, 02.06.2025)."
                )
            return

        if step == "end":
            try:
                d_to = datetime.strptime(text, "%d.%m.%Y").date()
                d_from = rp.get("date_from")
                if d_from and d_to < d_from:
                    await update.message.reply_text(
                        "Дата окончания раньше даты начала. Введите корректную дату окончания (ДД.ММ.ГГГГ)."
                    )
                    return
                rp["date_to"] = d_to
                rp["step"] = "done"
                context.user_data["redlamps_period"] = rp
                await update.message.reply_text(
                    f"Период сохранён: {d_from:%d.%m.%Y} — {d_to:%d.%m.%Y} (допуск ±{REDLAMPS_TOLERANCE_DAYS_DEFAULT} дней).\n"
                    "Теперь нажмите «📊 Сформировать BI-панель» в разделе «Красные лампочки»."
                )
            except Exception:
                await update.message.reply_text(
                    "Дата окончания в неверном формате. Введите ДД.ММ.ГГГГ (например, 12.06.2025)."
                )
            return
# Комментарий к доработке графика
    if context.user_data.get("awaiting_rework_comment"):
        info = context.user_data.pop("awaiting_rework_comment")
        version = info["version"]
        approver = info["approver"]
        comment = text
        update_schedule_approval_status(version, approver, "rework", comment)
        await update.message.reply_text(
            "Комментарий сохранён. График помечен как отправленный на доработку."
        )
        return

    # Ввод списка согласующих
    if context.user_data.get("awaiting_approvers_input"):
        info = context.user_data.pop("awaiting_approvers_input")
        version = info["version"]

        raw = text.replace(",", " ").split()
        approvers: List[str] = []
        for token in raw:
            token = token.strip()
            if not token:
                continue
            if not token.startswith("@"):
                token = "@" + token
            approvers.append(token)
        approvers = list(dict.fromkeys(approvers))

        if not approvers:
            await update.message.reply_text("Не найдено ни одного юзернейма.")
            return

        set_current_approvers_for_version(approvers, version)

        lines = [
            "График на новую неделю, необходимо согласовать.",
            f"Версия: {version}",
            "",
            "Согласующие:",
        ]
        for a in approvers:
            lines.append(f"• {a}")

        kb = InlineKeyboardMarkup(
            [
                [
                    InlineKeyboardButton(
                        f"✅ Согласовать ({a})", callback_data=f"schedule_approve:{a}"
                    ),
                    InlineKeyboardButton(
                        f"✏️ На доработку ({a})",
                        callback_data=f"schedule_rework:{a}",
                    ),
                ]
                for a in approvers
            ]
        )

        text_to_send = "\n".join(lines)

        await chat.send_message(text_to_send, reply_markup=kb)

        if SCHEDULE_NOTIFY_CHAT_ID is not None:
            try:
                await context.bot.send_message(
                    chat_id=SCHEDULE_NOTIFY_CHAT_ID,
                    text=text_to_send,
                    reply_markup=kb,
                )
            except Exception as e:
                log.error(
                    "Не удалось отправить уведомление в чат SCHEDULE_NOTIFY_CHAT_ID=%s: %s",
                    SCHEDULE_NOTIFY_CHAT_ID,
                    e,
                )

        await update.message.reply_text("Согласующие сохранены и уведомлены.")
        return

    # Поиск по номеру дела в замечаниях
    if context.user_data.get("awaiting_case_search"):
        context.user_data.pop("awaiting_case_search", None)
        case_no = text.strip()
        df = get_remarks_df_current()
        if df is None:
            await update.message.reply_text(
                "Не удалось открыть файл замечаний. Проверьте доступ к таблице."
            )
            return
        out_text = build_case_cards_text(df, case_no)
        await send_long_text(chat, out_text)
        return

    # Поиск по номеру дела в итоговых проверках
    if context.user_data.get("awaiting_final_case_search"):
        context.user_data.pop("awaiting_final_case_search", None)
        case_no = text.strip()
        df = get_final_checks_df()
        if df is None:
            await update.message.reply_text(
                "Не удалось открыть таблицу итоговых проверок."
            )
            return
        header = f"📋 Итоговые проверки по номеру дела: {case_no}"
        text_out = build_final_checks_text_filtered(
            df, case_no=case_no, header=header
        )
        await send_long_text(chat, text_out)
        await send_final_checks_xlsx_filtered(
            chat_id=chat.id, df=df, context=context, case_no=case_no
        )
        return

    low = text.lower()


    if low == "🚨 красные лампочки":
        kb = redlamps_menu_inline()
        has_file = "✅" if context.user_data.get("redlamps_file_bytes") else "❌"
        rp = context.user_data.get("redlamps_period") or {}
        has_period = "✅" if (rp.get("date_from") and rp.get("date_to")) else "❌"
        msg = (
            "🚨 Раздел «Красные лампочки»\n\n"
            "1) Загрузите Excel-файл (.xlsx)\n"
            "2) Выберите период (колонки K–L, допуск ±5 дней)\n"
            "3) Сформируйте BI-панель (счётчик актов/протоколов по дублям номера дела)\n\n"
            f"Файл: {has_file} | Период: {has_period}"
        )
        await update.message.reply_text(msg, reply_markup=kb)
        return
    if low == "📅 график".lower():
        settings = get_schedule_state()
        is_adm = is_admin(update.effective_user.id)
        msg = build_schedule_text(is_adm, settings)
        user_username = update.effective_user.username or ""
        user_tag = f"@{user_username}" if user_username else None
        kb = build_schedule_inline(is_adm, settings, user_tag=user_tag)
        msg_full = (
            "📅 Раздел «График выездов»\n\n"
            "• Смотреть текущий статус согласования\n"
            "• Обновить данные из общей таблицы\n"
            "• Скачать красиво оформленный Excel-файл\n\n"
            "Если вы входите в список согласующих, ниже будут кнопки "
            "«Согласовать» и «На доработку».\n\n"
            f"{msg}"
        )
        await update.message.reply_text(msg_full, reply_markup=kb)
        return

    if low == "📝 замечания".lower():
        kb = remarks_menu_inline()
        msg = (
            "📝 Раздел «Замечания»\n\n"
            "Здесь доступны:\n"
            "• 🔎 поиск по номеру дела (столбец I);\n"
            "• 🏗 ОНзС — выбор 1–12, список дел (Номер дела (I) + адрес) и отдельный просмотр неустранённых;\n"
            "• 📥 открыть общий файл таблицы.\n\n"
            "Выберите нужное действие:"
        )
        await update.message.reply_text(msg, reply_markup=kb)
        return

    if low in ("инспектор", "👮 инспектор"):
        kb = inspector_menu_inline()
        msg = (
            "👮‍♂️ Раздел «Инспектор»\n\n"
            "Здесь можно:\n"
            "• ➕ добавить выезд инспектора;\n"
            "• 📋 посмотреть последние выезды;\n"
            "• 📥 скачать отдельный Excel с выездами;\n"
            "• 🔄 обнулить список выездов (кнопка «Обновить»).\n\n"
            "Выберите действие кнопками ниже."
        )
        await update.message.reply_text(msg, reply_markup=kb)
        return

    if low == "📈 аналитика".lower():
        conn = get_db()
        c = conn.cursor()
        c.execute(
            """SELECT version, approver, status, comment, decided_at, requested_at
               FROM schedule_approvals
               ORDER BY version DESC, approver"""
        )
        rows = c.fetchall()
        conn.close()

        if not rows:
            await update.message.reply_text("Пока нет данных по согласованию графика.")
            return

        by_ver: Dict[int, List[sqlite3.Row]] = {}
        for r in rows:
            by_ver.setdefault(r["version"], []).append(r)

        lines: List[str] = ["📈 Аналитика по согласованию графика:", ""]

        for ver in sorted(by_ver.keys(), reverse=True):
            approvals = by_ver[ver]
            header = build_schedule_header(ver, approvals)
            lines.append("")
            lines.append(header + ":")
            for r in approvals:
                appr = r["approver"]
                status = r["status"] or "pending"
                decided = _format_dt(r["decided_at"])
                requested = _format_dt(r["requested_at"])
                comment = r["comment"] or ""

                if status == "pending":
                    lines.append(f"• {appr} — ожидает, запрошено {requested}")
                elif status == "approved":
                    lines.append(f"• {appr} — Согласовано {decided} ✅")
                elif status == "rework":
                    if comment:
                        lines.append(
                            f"• {appr} — На доработку {decided} (Комментарий: {comment})"
                        )
                    else:
                        lines.append(f"• {appr} — На доработку {decided}")

        await send_long_text(chat, "\n".join(lines))
        return

    if low == "итоговые проверки":
        # Каждый раз при входе в раздел обновляем локальный файл итоговых проверок
        ok = refresh_final_checks_local_file()
        if not ok:
            await update.message.reply_text(
                "Не удалось обновить файл итоговых проверок.\n"
                "Проверьте доступ к Google Sheets и переменную FINAL_CHECKS_SPREADSHEET_ID."
            )
            return

        # 1) Берём итоговые проверки из целевых листов (МКД/СОЦОБЪЕКТЫ/Остальное)
        df_final = get_final_checks_df_target_sheets()

        # 2) Берём таблицу замечаний (вторая таблица) — поиск по номеру дела (I)
        df_remarks = get_remarks_df_current()

        if df_final is None:
            await update.message.reply_text(
                "Не удалось открыть таблицу итоговых проверок (целевые листы)."
            )
            return

        if df_final.empty:
            await update.message.reply_text(
                "В целевых листах итоговых проверок нет строк с данными."
            )
            return

        if df_remarks is None or df_remarks.empty:
            await update.message.reply_text(
                "Не удалось открыть таблицу замечаний для проверки нарушений."
            )
            return

        # 3) Формируем «BI‑панель»: дела с датой начала (O) через 1–10 дней и со статусом «нет» в Q/R/Y/AD
        try:
            panel_text = build_final_checks_violations_bi_panel(
                df_final=df_final,
                df_remarks=df_remarks,
                days_min=1,
                days_max=10,
                report_day=local_now().date(),
            )
        except Exception as e:
            log.error("Ошибка формирования BI‑панели по итоговым проверкам: %s", e)
            await update.message.reply_text(
                "Не удалось сформировать список итоговых проверок с нарушениями."
            )
            return

        await send_long_text(chat, panel_text)
        return

    await update.message.reply_text(
        "Я вас не понял. Выберите пункт меню или нажмите /start.",
        reply_markup=main_menu(),
    )


# -------------------------------------------------
# DOCUMENT HANDLER
# -------------------------------------------------

async def document_handler(update: Update, context: ContextTypes.DEFAULT_TYPE):
    # 🚨 Красные лампочки — принимаем Excel-файл только когда пользователь нажал «Загрузить файл»
    if context.user_data.get("awaiting_redlamps_upload"):
        doc = update.message.document
        if not doc:
            await update.message.reply_text("Не удалось получить файл. Повторите отправку.")
            return

        fname = (doc.file_name or "").lower()
        if not fname.endswith(".xlsx"):
            await update.message.reply_text("Нужен файл Excel в формате .xlsx. Отправьте корректный файл.")
            return

        try:
            tg_file = await context.bot.get_file(doc.file_id)
            bio = BytesIO()
            await tg_file.download_to_memory(out=bio)
            bio.seek(0)
            context.user_data["redlamps_file_bytes"] = bio.getvalue()
            context.user_data["redlamps_file_name"] = doc.file_name or "upload.xlsx"
            context.user_data.pop("awaiting_redlamps_upload", None)

            rp = context.user_data.get("redlamps_period") or {}
            has_period = "✅" if (rp.get("date_from") and rp.get("date_to")) else "❌"

            await update.message.reply_text(
                f"✅ Файл загружен: {doc.file_name}\n"
                f"Период: {has_period}\n\n"
                "Далее: задайте период кнопкой «📅 Выбрать период (K–L)» и нажмите «📊 Сформировать BI-панель».",
                reply_markup=redlamps_menu_inline(),
            )
            return
        except Exception as e:
            log.error("REDLAMPS: ошибка загрузки файла: %s", e)
            await update.message.reply_text("Ошибка при загрузке файла. Попробуйте ещё раз.")
            return

    # по умолчанию загрузка файлов в боте отключена
    await update.message.reply_text(
        "Загрузка файлов через бота отключена. Используйте общую Google-таблицу."
    )




# -------------------------------------------------
# VOICE (SpeechKit STT) -> ассистент
# -------------------------------------------------
async def voice_router(update: Update, context: ContextTypes.DEFAULT_TYPE):
    if not ENABLE_ASSISTANT:
        await update.message.reply_text("Голосовой ассистент отключён администратором.")
        return
    try:
        file = await context.bot.get_file(update.message.voice.file_id)
        ogg = await file.download_as_bytearray()
        text_q = yandex_speech_to_text(bytes(ogg))
        # включаем режим автоматически
        context.user_data["assistant_mode"] = True
        await assistant_answer(update.message.chat, context, text_q, recognized_from_voice=True)
    except Exception as e:
        await update.message.reply_text(f"Не удалось обработать голосовое: {e}")


# -------------------------------------------------
# START / HELP
# -------------------------------------------------
async def start(update: Update, context: ContextTypes.DEFAULT_TYPE):
    msg = (
        "Добро пожаловать в бота отдела СОТ.\n\n"
        "Основные разделы:\n"
        "• 📅 График — согласование графика выездов\n"
        "• 📝 Замечания — поиск по номеру дела, ОНзС и статусы «нет»\n"
        "• Инспектор — выезды инспектора\n"
        "• Итоговые проверки — перечень итоговых проверок по отдельной таблице\n"
        "• 🚨 Красные лампочки — загрузка Excel и BI-панель по актам/протоколам\n"
        "• 📈 Аналитика — история согласований\n\n"
        "Выберите раздел с помощью кнопок ниже."
    )
    await update.message.reply_text(msg, reply_markup=main_menu())


async def help_command(update: Update, context: ContextTypes.DEFAULT_TYPE):
    msg = (
        "Справка по боту СОТ:\n\n"
        "📅 График — показать статус согласования, обновить, скачать Excel.\n"
        "📝 Замечания — поиск по номеру дела (I), работа с ОНзС и просмотр статусов «нет».\n"
        "Инспектор — добавление и выгрузка выездов инспектора.\n"
        "Итоговые проверки — список и выгрузка итоговых проверок за период или по делу.\n"
        "📈 Аналитика — история согласований по версиям графика.\n"
    )
    await update.message.reply_text(msg, reply_markup=main_menu())


# -------------------------------------------------
# MAIN
# -------------------------------------------------
async def cb_file_action(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    q = update.callback_query
    if not q:
        return
    await q.answer()
    data = q.data or ""
    try:
        _, action, token = data.split(":", 2)
    except Exception:
        return
    meta = FILE_TOKEN_MAP.get(token)
    if not meta:
        await q.edit_message_text("Ссылка устарела. Повторите запрос.")
        return
    url = meta.get("url", "")
    if not url:
        await q.edit_message_text("Ссылка не найдена.")
        return

    if action == "open":
        await q.edit_message_text(f"Открыть ссылку: {url}", disable_web_page_preview=True)
        return

    try:
        content, fname, mime = download_external_file(url)
    except Exception as e:
        await q.edit_message_text(f"Не удалось скачать файл: {e}\n\nСсылка: {url}", disable_web_page_preview=True)
        return

    if action == "download":
        await context.bot.send_document(
            chat_id=q.message.chat_id,
            document=content,
            filename=fname,
            caption=f"Файл: {fname}\nИсточник: {url}"
        )
        return

    if action == "analyze":
        analysis = analyze_file_bytes(content, fname, mime)
        await context.bot.send_message(
            chat_id=q.message.chat_id,
            text=f"🔎 Анализ: {fname}\n\n{analysis}\n\nИсточник: {url}",
            disable_web_page_preview=True
        )
        await context.bot.send_document(
            chat_id=q.message.chat_id,
            document=content,
            filename=fname,
            caption=f"Файл для просмотра: {fname}"
        )
        return


# =================================================
# 🧪 ТЗ для ЦНИЛ — Чек-лист (мастер)
# =================================================
CNIL_MENU_LABEL = "🧪 ТЗ для ЦНИЛ"
CNIL_CASE_RE = re.compile(r"\b(\d{2})[-\s]?(\d{2})[-\s]?(\d{6})\b")

# Начальники территориальных отделов Госстройнадзора МО (по первым 2 цифрам номера дела)
CNIL_HEAD_BY_PREFIX = {
    "01": "Герасименко Д.А.",
    "02": "Кузьмичев Е.М.",
    "03": "Маркелов А.С.",
    "04": "Гутнов З.В.",
    "05": "Арменакян Г.Б.",
    "06": "Краснов В.А.",
    "07": "Денисов Д.М.",
    "08": "Романова Л.П.",
    "09": "Садоян Д.Т.",
    "10": "Павлов А.В.",
    "11": "Ефимов Р.С.",
    "12": "Нестеров И.М.",
}


def cnil_head_for_case(case_no: str) -> str:
    """Возвращает ФИО начальника по первым 2 цифрам номера дела."""
    try:
        prefix = (case_no or "").strip().split("-")[0].zfill(2)
    except Exception:
        prefix = ""
    return CNIL_HEAD_BY_PREFIX.get(prefix, "Гутнов З.В.")


def _docx_replace_paragraph_text(paragraph, new_text: str) -> None:
    """Заменяет текст абзаца целиком (теряет разноформатные run'ы внутри абзаца, но сохраняет стиль абзаца)."""
    # python-docx: paragraph.text is writable
    paragraph.text = new_text


def _docx_regex_replace(doc, pattern: re.Pattern, repl_func) -> int:
    """Применяет regex-замену ко всем абзацам (и таблицам) документа. Возвращает число заменённых абзацев."""
    changed = 0

    def handle_paragraph(paragraph):
        nonlocal changed
        txt = paragraph.text
        if not txt:
            return
        new_txt = pattern.sub(repl_func, txt)
        if new_txt != txt:
            _docx_replace_paragraph_text(paragraph, new_txt)
            changed += 1

    for para in doc.paragraphs:
        handle_paragraph(para)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    handle_paragraph(para)

    return changed


def cnil_generate_t3_docx(row: dict) -> str:
    """Generate filled TEST_T3 .docx based on template.

    The template contains a 3-column table (Индекс работы / Состав работ / Краткая информация...).
    We rebuild the table rows exactly to the number of selected works so that it automatically
    grows/shrinks depending on the user's selection.

    Expected row keys (best-effort):
      - case_no / Номер дела
      - works / Работы (comma-separated or semicolon-separated string)
      - head_name (optional) - territorial head FIO to be inserted
    """
    from docx import Document
    from docx.shared import Pt
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    # Требование: весь добавленный текст должен быть Times New Roman 12
    FONT_NAME = "Times New Roman"
    FONT_SIZE = Pt(12)

    def _set_run_font(run):
        """Принудительно выставляет шрифт и кегль для run (включая EastAsia)."""
        try:
            run.font.name = FONT_NAME
            run.font.size = FONT_SIZE
            rpr = run._element.get_or_add_rPr()
            rfonts = rpr.find(qn('w:rFonts'))
            if rfonts is None:
                rfonts = OxmlElement('w:rFonts')
                rpr.append(rfonts)
            rfonts.set(qn('w:ascii'), FONT_NAME)
            rfonts.set(qn('w:hAnsi'), FONT_NAME)
            rfonts.set(qn('w:cs'), FONT_NAME)
            rfonts.set(qn('w:eastAsia'), FONT_NAME)
        except Exception:
            # На случай несовместимых версий python-docx: не падаем
            pass

    def _set_paragraph_font(p):
        """Выставляет Times New Roman 12 для всех runs параграфа (создаёт run при необходимости)."""
        if not p.runs:
            _set_run_font(p.add_run(""))
        for r in p.runs:
            _set_run_font(r)

    def _get_case_no(r: dict) -> str:
        for k in ('case_no', 'Номер дела', 'номер дела', 'case', 'дело'):
            v = r.get(k)
            if v:
                return str(v).strip()
        return ''

    def _split_works(val) -> list[str]:
        if not val:
            return []
        s = str(val).strip()
        # common separators in this bot
        parts = re.split(r'[\n;]+', s)
        out = [p.strip() for p in parts if p and p.strip()]
        return out

    def _parse_work(item: str) -> tuple[str, str]:
        """Return (index, text). If no explicit index, index is empty."""
        m = re.match(r'^\s*([0-9]+(?:\.[0-9]+)*)\s+(.*)$', item)
        if m:
            return m.group(1), m.group(2).strip()
        return '', item.strip()

    def _set_table_autofit(table):
        # python-docx has limited support; set both flags + XML tblLayout
        try:
            table.autofit = True
        except Exception:
            pass
        tbl = table._tbl
        tblPr = tbl.tblPr
        tblLayout = tblPr.find(qn('w:tblLayout'))
        if tblLayout is None:
            tblLayout = OxmlElement('w:tblLayout')
            tblPr.append(tblLayout)
        tblLayout.set(qn('w:type'), 'autofit')

    def _set_table_borders(table):
        # Ensure visible borders for the whole table (Word can drop gridlines on rebuilt tables)
        try:
            tbl = table._tbl
            tblPr = tbl.tblPr
            # Remove existing tblBorders (if any)
            for el in list(tblPr):
                if el.tag.endswith('tblBorders'):
                    tblPr.remove(el)
            tblBorders = OxmlElement('w:tblBorders')
            for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
                e = OxmlElement(f'w:{edge}')
                e.set(qn('w:val'), 'single')
                e.set(qn('w:sz'), '8')  # 1/8 pt units
                e.set(qn('w:space'), '0')
                e.set(qn('w:color'), '000000')
                tblBorders.append(e)
            tblPr.append(tblBorders)
        except Exception:
            pass

        # Also set borders per-cell for maximum compatibility
        try:
            for row in table.rows:
                for cell in row.cells:
                    tcPr = cell._tc.get_or_add_tcPr()
                    # Remove existing tcBorders (if any)
                    for el in list(tcPr):
                        if el.tag.endswith('tcBorders'):
                            tcPr.remove(el)
                    tcBorders = OxmlElement('w:tcBorders')
                    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
                        e = OxmlElement(f'w:{edge}')
                        e.set(qn('w:val'), 'single')
                        e.set(qn('w:sz'), '8')
                        e.set(qn('w:space'), '0')
                        e.set(qn('w:color'), '000000')
                        tcBorders.append(e)
                    tcPr.append(tcBorders)
        except Exception:
            pass


    # inputs
    case_no = _get_case_no(row)
    works_items = _split_works(row.get('works') or row.get('Работы') or row.get('works_text'))
    head_name = (row.get('head_name') or row.get('Начальник') or row.get('Начальник территориального отдела') or '').strip()
    if (not head_name) and case_no:
        head_name = cnil_head_for_case(case_no)


    tpl = CNIL_T3_TEMPLATE
    if not os.path.exists(tpl):
        raise FileNotFoundError(f"CNIL template not found: {tpl}")

    doc = Document(tpl)

    # Ensure newly created/modified text is in Times New Roman 12.
    # Setting Normal helps runs created by python-docx, while we also explicitly
    # enforce on paragraphs/cells we touch.
    try:
        normal_font = doc.styles["Normal"].font
        normal_font.name = FONT_NAME
        normal_font.size = FONT_SIZE
    except Exception:
        pass

    def _apply_font_to_paragraph(paragraph):
        if paragraph is None:
            return
        if not paragraph.runs:
            r = paragraph.add_run("")
            _set_run_font(r)
            return
        for r in paragraph.runs:
            _set_run_font(r)

    # Replace the 'Код объекта' line robustly (do not rely on a hardcoded code)
    if case_no:
        for p in doc.paragraphs:
            if 'Код объекта' in p.text:
                # Normalize: 'Код объекта: № XX-XX-XXXXXX.'
                p.text = re.sub(r'Код объекта\s*:\s*№\s*[^\.\n]+', f'Код объекта: № {case_no}', p.text)
                _apply_font_to_paragraph(p)

    # Replace head FIO if provided (keep template default otherwise)
    if head_name:
        for p in doc.paragraphs:
            if 'Начальник территориального отдела' in p.text and 'Госстройнадзора МО' in p.text:
                # replace everything after underscore/line to FIO
                p.text = re.sub(r'(Начальник территориального отдела\s*Госстройнадзора\s*МО\s*[^_\n]*_\s*)([^\n]+)$', r'\1' + head_name, p.text)
                _apply_font_to_paragraph(p)

                p.text = re.sub(r'(Начальник территориального отдела\s*Госстройнадзора\s*МО\s*.*?\s)\S+\s*[А-ЯЁ]\.[А-ЯЁ]\.$', r'\1' + head_name, p.text)
                _apply_font_to_paragraph(p)

    # Rebuild the works table to match number of works
    if works_items:
        # heuristic: choose the first 4-column table whose header contains 'Индекс', 'Состав' and 'Отметка'
        target = None
        for t in doc.tables:
            try:
                if (
                    len(t.columns) >= 4
                    and t.cell(0, 0).text.strip().startswith('Индекс')
                    and 'Состав' in t.cell(0, 1).text
                    and ('Отметка' in t.cell(0, 2).text or 'ОСИ' in t.cell(0, 2).text)
                ):
                    target = t
                    break
            except Exception:
                continue

        if target is not None:
            _set_table_autofit(target)
            _set_table_borders(target)

            # keep header row (row 0), remove all other rows
            while len(target.rows) > 1:
                target._tbl.remove(target.rows[1]._tr)

            wm = row.get("work_marks") or {}

            for item in works_items:
                idx, txt = _parse_work(item)
                r = target.add_row()
                # column 0: index
                r.cells[0].text = idx
                # column 1: composition
                r.cells[1].text = ('- ' + txt) if txt and not txt.startswith('-') else txt
                # column 2: marks per work
                if isinstance(wm, dict):
                    r.cells[2].text = wm.get(item, '') or wm.get(txt, '')
                else:
                    r.cells[2].text = ''
                # column 3: short info
                r.cells[3].text = ''

                # Enforce font on inserted cell text
                for c in r.cells:
                    for p in c.paragraphs:
                        _set_paragraph_font(p)

    # Output path: do not depend on a global DATA_DIR (it may be None in some deployments)
    try:
        out_dir = str(cnil_data_dir())
    except Exception:
        out_dir = os.getenv("DATA_DIR") or ("/data" if os.path.isdir("/data") else "data")
    os.makedirs(out_dir, exist_ok=True)
    safe_case = case_no if case_no else "NOCASE"
    out_path = os.path.join(out_dir, f"T3_{safe_case}.docx")
    doc.save(out_path)
    return out_path

def cnil_data_dir() -> Path:
    # Railway volume обычно монтируется в /data
    env_dir = os.getenv("DATA_DIR")
    if env_dir:
        p = Path(env_dir)
    else:
        p = Path("/data") if Path("/data").exists() else Path("data")
    p.mkdir(parents=True, exist_ok=True)
    return p

def cnil_results_path() -> Path:
    return cnil_data_dir() / "cnil_results.xlsx"


def cnil_append_to_gsheet(row: dict) -> str:
    """
    Append one CNIL record to Google Sheets.

    Config:
      - CNIL_GSHEET_ID: Spreadsheet ID (defaults to the provided sheet)
      - CNIL_SHEET_NAME: Worksheet title (default: "ТЗ для ЦНИЛ")
      - Credentials (one of):
          * GOOGLE_SERVICE_ACCOUNT_JSON : service account json as a string
          * GOOGLE_APPLICATION_CREDENTIALS : path to json file
          * GOOGLE_SERVICE_ACCOUNT_FILE : path to json file
    """
    sheet_id = os.getenv("CNIL_GSHEET_ID", "10sIT5I1WIkg2YzNHUpQOgeW2tN-PeHhZbogzp3G468s")
    ws_name = os.getenv("CNIL_SHEET_NAME", "ТЗ для ЦНИЛ")

    sa_json = os.getenv("GOOGLE_SERVICE_ACCOUNT_JSON")
    sa_path = os.getenv("GOOGLE_APPLICATION_CREDENTIALS") or os.getenv("GOOGLE_SERVICE_ACCOUNT_FILE")

    try:
        import gspread
        from google.oauth2.service_account import Credentials
    except Exception as e:
        raise RuntimeError(f"gspread/google-auth не установлены: {e}")

    if sa_json:
        try:
            info = json.loads(sa_json)
        except Exception as e:
            raise RuntimeError(f"GOOGLE_SERVICE_ACCOUNT_JSON не является корректным JSON: {e}")
        creds = Credentials.from_service_account_info(info, scopes=[
            "https://www.googleapis.com/auth/spreadsheets",
            "https://www.googleapis.com/auth/drive",
        ])
    elif sa_path:
        p = Path(sa_path)
        if not p.exists():
            raise RuntimeError(f"Файл ключа service account не найден: {sa_path}")
        creds = Credentials.from_service_account_file(str(p), scopes=[
            "https://www.googleapis.com/auth/spreadsheets",
            "https://www.googleapis.com/auth/drive",
        ])
    else:
        raise RuntimeError("Не заданы креды Google: задайте GOOGLE_SERVICE_ACCOUNT_JSON или GOOGLE_APPLICATION_CREDENTIALS/GOOGLE_SERVICE_ACCOUNT_FILE")

    gc = gspread.authorize(creds)
    sh = gc.open_by_key(sheet_id)
    try:
        ws = sh.worksheet(ws_name)
    except Exception:
        ws = sh.add_worksheet(title=ws_name, rows=1000, cols=20)

    header = ["timestamp", "user_id", "username", "case_no", "stage", "element", "marks_axes", "defects", "works"]
    # if sheet is empty, put header
    try:
        first = ws.row_values(1)
    except Exception:
        first = []
    if not first:
        ws.append_row(header, value_input_option="USER_ENTERED")

    values = [row.get(k, "") for k in header]
    ws.append_row(values, value_input_option="USER_ENTERED")
    return getattr(sh, "url", f"https://docs.google.com/spreadsheets/d/{sheet_id}/edit")


def cnil_save_row(row: dict) -> dict:
    """
    Save CNIL record:
      1) Google Sheets (if creds provided)
      2) Local Excel fallback (/data/cnil_results.xlsx)
    Returns:
      {"backend":"gsheet","url": "..."} or {"backend":"excel","path": "..."}
    """
    prefer_gs = os.getenv("CNIL_SAVE_TO_GSHEETS", "1") not in ("0", "false", "False", "no", "NO")
    if prefer_gs:
        try:
            url = cnil_append_to_gsheet(row)
            return {"backend": "gsheet", "url": url}
        except Exception:
            logging.exception("[CNIL] failed to save to Google Sheets, falling back to local Excel")

    out_path = cnil_results_path()
    if out_path.exists():
        df_old = pd.read_excel(out_path)
        df_new = pd.concat([df_old, pd.DataFrame([row])], ignore_index=True)
    else:
        df_new = pd.DataFrame([row])
    df_new.to_excel(out_path, index=False)
    return {"backend": "excel", "path": str(out_path)}


def cnil_ensure_results_file(path: Path) -> None:
    """Создаёт файл результатов ЦНИЛ с заголовками, если он отсутствует."""
    if path.exists():
        return
    df = pd.DataFrame(
        columns=[
            "ts",
            "user_id",
            "username",
            "full_name",
            "case_no",
            "stage",
            "element",
            "marks_axes",
            "defects",
            "works",
        ]
    )
    df.to_excel(path, index=False)


def cnil_pretty_export(src: Path, dst: Path) -> None:
    """Делает 'красивую' выгрузку: шапка, фильтры, закрепление, ширины, перенос строк."""
    from openpyxl import load_workbook
    from openpyxl.styles import Font, Alignment, PatternFill

    wb = load_workbook(src)
    ws = wb.active

    # Заголовок
    header_fill = PatternFill("solid", fgColor="F2F2F2")
    header_font = Font(bold=True)
    header_align = Alignment(horizontal="center", vertical="center", wrap_text=True)
    for cell in ws[1]:
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = header_align

    ws.freeze_panes = "A2"
    ws.auto_filter.ref = ws.dimensions

    # Перенос строк для 'defects' и 'works' + выравнивание по верхнему краю
    wrap_top = Alignment(vertical="top", wrap_text=True)
    max_row = ws.max_row
    max_col = ws.max_column
    for r in range(2, max_row + 1):
        for c in range(1, max_col + 1):
            ws.cell(row=r, column=c).alignment = wrap_top

    # Авто-ширины (с ограничением)
    for col_cells in ws.columns:
        col_letter = col_cells[0].column_letter
        max_len = 0
        for cell in col_cells:
            v = cell.value
            if v is None:
                continue
            s = str(v)
            if len(s) > max_len:
                max_len = len(s)
        ws.column_dimensions[col_letter].width = max(10, min(60, max_len + 2))

    wb.save(dst)


async def cnil_send_results_excel(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    """Отправляет файл результатов ЦНИЛ пользователю (после проверки пароля)."""
    src = Path(DATA_DIR) / "cnil_results.xlsx"
    cnil_ensure_results_file(src)

    # Формируем 'красивую' копию для скачивания (чтобы не ломать файл, который дописываем)
    dst = Path(DATA_DIR) / "cnil_results_pretty.xlsx"
    try:
        cnil_pretty_export(src, dst)
        send_path = dst
    except Exception:
        # если форматирование не удалось — отправим исходник
        send_path = src

    await update.message.reply_document(
        document=open(send_path, "rb"),
        filename="cnil_results.xlsx",
        caption=(
            "📄 Результаты ЦНИЛ\n"
            "Структурированная таблица с заполненными заявками."
        ),
        reply_markup=cnil_menu_keyboard(),
    )


def cnil_find_excel() -> Optional[Path]:
    """
    Ищем файл справочника в типовых местах.
    """
    candidates = [
        Path(os.getenv("CNIL_XLSX", "")) if os.getenv("CNIL_XLSX") else None,
        Path("/data/Чек-лист ГСН-ЦНИЛ.xlsx"),
        Path("/app/Чек-лист ГСН-ЦНИЛ.xlsx"),
        Path("Чек-лист ГСН-ЦНИЛ.xlsx"),
    ]
    for c in candidates:
        if c and c.exists():
            return c
    return None

def cnil_load_catalog() -> dict:
    """
    Строим иерархию из Excel (ваш чек-лист):
    stage -> element -> defect -> list(works)

    ВАЖНО: в исходном файле заголовки и порядок колонок фиксированы (5 колонок),
    поэтому мы используем позиционное сопоставление, чтобы исключить ошибки
    распознавания по ключевым словам.
    Также «Цоколь» по логике пользователя относится к разделу «Фундамент»,
    поэтому маппим stage='Цоколь' -> stage='Фундамент'.
    """
    xlsx = cnil_find_excel()
    if not xlsx:
        return {}

    try:
        df = pd.read_excel(xlsx, sheet_name=0)
    except Exception:
        logging.exception("[CNIL] failed to read excel")
        return {}

    # Ожидаемые 5 колонок (как в файле пользователя)
    if df.shape[1] < 5:
        logging.error("[CNIL] unexpected columns count: %s", df.shape[1])
        return {}

    # Позиционное сопоставление колонок (надежнее, чем fuzzy-поиск)
    col_stage = df.columns[0]
    col_elem  = df.columns[1]
    col_def   = df.columns[3]
    col_work  = df.columns[4]

    def clean(x) -> str:
        if x is None:
            return ""
        s = str(x).strip()
        if not s or s.lower() in ("nan", "none"):
            return ""
        # отбрасываем служебные номера 1/2/3/4/5 в верхних строках
        if re.fullmatch(r"\d+(?:\.\d+)?", s):
            return ""
        return s

    # ffill для многострочных блоков
    df[col_stage] = df[col_stage].ffill()
    df[col_elem]  = df[col_elem].ffill()
    df[col_def]   = df[col_def].ffill()

    cat: dict = {}
    for _, r in df.iterrows():
        stage_raw = clean(r.get(col_stage))
        elem  = clean(r.get(col_elem))
        defect = clean(r.get(col_def))
        work  = clean(r.get(col_work))

        if not (stage_raw and elem and defect and work):
            continue

        # По ТЗ: «Цоколь» относится к «Фундамент»
        stage = "Фундамент" if stage_raw == "Цоколь" else stage_raw

        cat.setdefault(stage, {}).setdefault(elem, {}).setdefault(defect, set()).add(work)

    # set -> sorted list (стабильный порядок)
    out: dict = {}
    for st, elems in cat.items():
        out[st] = {}
        for el, defects in elems.items():
            out[st][el] = {}
            for dfct, works in defects.items():
                out[st][el][dfct] = sorted(works)
    return out

CNIL_CATALOG = cnil_load_catalog()

def cnil_norm_case(text: str) -> Optional[str]:
    m = CNIL_CASE_RE.search(text or "")
    if not m:
        return None
    return f"{m.group(1)}-{m.group(2)}-{m.group(3)}"

def cnil_kb_list(items: list[str], prefix: str, page: int = 0, page_size: int = 8) -> InlineKeyboardMarkup:
    total = len(items)
    pages = max(1, (total + page_size - 1) // page_size)
    page = max(0, min(page, pages - 1))
    start = page * page_size
    chunk = items[start:start + page_size]

    rows = []
    for i, label in enumerate(chunk):
        rows.append([InlineKeyboardButton(label, callback_data=f"cnil:{prefix}:pick:{start+i}")])

    nav = []
    if pages > 1:
        if page > 0:
            nav.append(InlineKeyboardButton("⬅️", callback_data=f"cnil:{prefix}:page:{page-1}"))
        nav.append(InlineKeyboardButton(f"{page+1}/{pages}", callback_data="cnil:noop"))
        if page < pages - 1:
            nav.append(InlineKeyboardButton("➡️", callback_data=f"cnil:{prefix}:page:{page+1}"))
        rows.append(nav)

    rows.append([InlineKeyboardButton("✖️ Отмена", callback_data="cnil:cancel")])
    return InlineKeyboardMarkup(rows)

def cnil_kb_multi(items: list[str], selected: set[int], prefix: str, page: int = 0, page_size: int = 8) -> InlineKeyboardMarkup:
    total = len(items)
    pages = max(1, (total + page_size - 1) // page_size)
    page = max(0, min(page, pages - 1))
    start = page * page_size
    chunk = items[start:start + page_size]

    rows = []
    for i, label in enumerate(chunk):
        idx = start + i
        mark = "✅ " if idx in selected else "⬜️ "
        rows.append([InlineKeyboardButton(mark + label, callback_data=f"cnil:{prefix}:toggle:{idx}")])

    nav = []
    if pages > 1:
        if page > 0:
            nav.append(InlineKeyboardButton("⬅️", callback_data=f"cnil:{prefix}:page:{page-1}"))
        nav.append(InlineKeyboardButton(f"{page+1}/{pages}", callback_data="cnil:noop"))
        if page < pages - 1:
            nav.append(InlineKeyboardButton("➡️", callback_data=f"cnil:{prefix}:page:{page+1}"))
        rows.append(nav)

    rows.append([
        InlineKeyboardButton("🧹 Сброс", callback_data=f"cnil:{prefix}:reset"),
        InlineKeyboardButton("Далее ➡️", callback_data=f"cnil:{prefix}:next"),
    ])
    rows.append([InlineKeyboardButton("⬅️ Назад", callback_data=f"cnil:{prefix}:back")])
    rows.append([InlineKeyboardButton("✖️ Отмена", callback_data="cnil:cancel")])
    return InlineKeyboardMarkup(rows)

def cnil_menu_keyboard() -> ReplyKeyboardMarkup:
    return ReplyKeyboardMarkup(
        [
            ["📝 Заполнить форму"],
            ["⬇️ Скачать таблицу"],
            ["🔐 Изменить пароль скачивания"],
            ["⬅️ Назад"],
        ],
        resize_keyboard=True,
    )


async def cnil_start(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Открыть раздел 🧪 ТЗ для ЦНИЛ (меню действий)."""
    context.user_data.pop("cnil", None)
    context.user_data.pop("cnil_download_wait", None)
    context.user_data.pop("cnil_change_step", None)
    await update.message.reply_text(
        "🧪 *ТЗ для ЦНИЛ* — выберите действие:",
        parse_mode="Markdown",
        reply_markup=cnil_menu_keyboard(),
    )


async def cnil_start_form(update: Update, context: ContextTypes.DEFAULT_TYPE):
    # каталог может не загрузиться — тогда предупредим
    if not CNIL_CATALOG:
        await update.message.reply_text(
            "⚠️ Не удалось загрузить справочник ЦНИЛ из Excel.\n"
            "Убедитесь, что файл «Чек-лист ГСН-ЦНИЛ.xlsx» находится в проекте (корень или /data).\n\n"
            "Пока мастер недоступен.",
            reply_markup=main_menu(),
        )
        return

    context.user_data["cnil"] = {
        "step": "case",
        "case_no": None,
        "stage": None,
        "element": None,
        "defects_selected": set(),
        "works_selected": set(),
        "work_marks": {},
        "work_marks_order": [],
        "work_marks_idx": 0,
        "page": 0,
        "items": {},
    }
    await update.message.reply_text(
        "1/5. Введите номер дела (формат 00-00-000000):",
        reply_markup=main_menu(),  # главное меню НЕ убираем
    )

async def cnil_text_step(update: Update, context: ContextTypes.DEFAULT_TYPE):
    st = context.user_data.get("cnil") or {}
    step = st.get("step")
    msg = (update.message.text or "").strip()

    if step == "case":
        case_no = cnil_norm_case(msg)
        if not case_no:
            await update.message.reply_text("Неверный формат. Введите номер дела: 00-00-000000")
            return
        st["case_no"] = case_no
        st["step"] = "stage"
        st["page"] = 0
        context.user_data["cnil"] = st

        stages = sorted(CNIL_CATALOG.keys())
        st["items"]["stage"] = stages
        kb = cnil_kb_list(stages, prefix="stage", page=0)
        await update.message.reply_text("2/5. Выберите этап строительства:", reply_markup=kb)
        return

    if step == "work_marks":
        # Ввод "Отметка ОСИ, высотные отметки" для каждой выбранной работы (можно пропустить)
        order = st.get("work_marks_order") or []
        idx = int(st.get("work_marks_idx") or 0)

        if not order:
            st["step"] = "works"
            context.user_data["cnil"] = st
            await update.message.reply_text("Выберите виды работ кнопками под сообщением.")
            return

        if idx < 0:
            idx = 0
        if idx >= len(order):
            idx = len(order)

        if idx < len(order):
            work_name = order[idx]
            val = (update.message.text or "").strip()
            if val in ("-", "—", ""):
                val = ""
            st.setdefault("work_marks", {})[work_name] = val
            idx += 1
            st["work_marks_idx"] = idx
            context.user_data["cnil"] = st

            if idx < len(order):
                next_work = order[idx]
                await update.message.reply_text(
                    f"Работа: {next_work}\nВведите «Отметка ОСИ, высотные отметки» (можно пропустить):",
                    reply_markup=InlineKeyboardMarkup([[InlineKeyboardButton("✖️ Отмена", callback_data="cnil:cancel")]]),
                )
                return

        # --- Все отметки собраны: формируем итоговую запись и документ ---
        user = update.effective_user
        stage = st.get("stage")
        element = st.get("element")
        case_no = st.get("case_no")

        defects_all = st.get("items", {}).get("defects", [])
        defects_sel = st.get("defects_selected") or set()
        defects = [defects_all[i] for i in sorted(defects_sel)] if defects_all else []

        works_all = st.get("items", {}).get("works", [])
        works_sel = st.get("works_selected") or set()
        works = [works_all[i] for i in sorted(works_sel)] if works_all else order

        wm = st.get("work_marks") or {}
        marks_axes = "\n".join([f"{w}: {wm.get(w, '')}".rstrip() for w in works]).strip()

        row = {
            "ts": datetime.utcnow().strftime("%Y-%m-%d %H:%M:%S"),
            "user_id": user.id if user else "",
            "username": (user.username or "") if user else "",
            "full_name": f"{(user.first_name or '')} {(user.last_name or '')}".strip() if user else "",
            "case_no": case_no,
            "stage": stage,
            "element": element,
            "marks_axes": marks_axes,
            "defects": "; ".join(defects),
            "works": "; ".join(works),
            "work_marks": wm,
        }

        try:
            save_info = cnil_save_row(row)
        except Exception as e:
            await update.message.reply_text(f"❌ Не удалось сохранить результат: {e}")
            context.user_data.pop("cnil", None)
            return

        target_line = (
            f"Google Sheets: {save_info.get('url')}" if save_info.get("backend") == "gsheet" else f"Файл (fallback): {save_info.get('path')}"
        )

        summary = (
            "✅ Сохранено.\n\n"
            f"Номер дела: {case_no}\n"
            f"Этап: {stage}\n"
            f"Элемент: {element}\n"
            "Отметка ОСИ, высотные отметки:\n"
            f"{marks_axes or '-'}\n"
            f"Дефекты: {', '.join(defects) if defects else '-'}\n"
            f"Работы: {', '.join(works) if works else '-'}\n\n"
            f"{target_line}"
        )

        await update.message.reply_text(summary)

        try:
            docx_path = cnil_generate_t3_docx(row)
            with open(docx_path, "rb") as f:
                await context.bot.send_document(
                    chat_id=update.effective_chat.id,
                    document=InputFile(f, filename=os.path.basename(docx_path)),
                    caption="ТЗ (TEST_T3) — заполненный документ",
                )
        except Exception as e:
            await context.bot.send_message(chat_id=update.effective_chat.id, text=f"⚠️ Документ TEST_T3 не сформирован: {e}")

        context.user_data.pop("cnil", None)
        return


        return

    # если пользователь пишет во время выбора кнопками — мягко подскажем
    if step in ("stage", "element", "defects", "works"):
        await update.message.reply_text("Используйте кнопки под сообщением для выбора.")
        return

async def cnil_callback(update: Update, context: ContextTypes.DEFAULT_TYPE):
    q = update.callback_query
    data = q.data or ""
    st = context.user_data.get("cnil") or {}
    step = st.get("step")

    if data == "cnil:noop":
        await q.answer()
        return

    if data == "cnil:cancel":
        context.user_data.pop("cnil", None)
        await q.edit_message_text("Мастер «ТЗ для ЦНИЛ» отменён.", reply_markup=None)
        return

    parts = data.split(":")
    # cnil:<scope>:<action>:<value?>
    if len(parts) < 3:
        await q.answer()
        return

    scope = parts[1]
    action = parts[2]
    value = parts[3] if len(parts) > 3 else None

    # ---------- ЭТАП ----------
    if scope == "stage":
        items = st.get("items", {}).get("stage", [])
        if action == "page" and value is not None:
            page = int(value)
            kb = cnil_kb_list(items, prefix="stage", page=page)
            await q.edit_message_reply_markup(reply_markup=kb)
            return
        if action == "pick" and value is not None:
            idx = int(value)
            if idx < 0 or idx >= len(items):
                await q.answer()
                return
            stage = items[idx]
            st["stage"] = stage
            st["step"] = "element"
            st["page"] = 0

            elements = sorted(CNIL_CATALOG.get(stage, {}).keys())
            st["items"]["element"] = elements
            context.user_data["cnil"] = st

            kb = cnil_kb_list(elements, prefix="element", page=0)
            await q.edit_message_text(f"Выбран этап: {stage}\n\n3/5. Выберите конструктивный элемент:", reply_markup=kb)
            return

    # ---------- ЭЛЕМЕНТ ----------
    if scope == "element":
        items = st.get("items", {}).get("element", [])
        if action == "page" and value is not None:
            page = int(value)
            kb = cnil_kb_list(items, prefix="element", page=page)
            await q.edit_message_reply_markup(reply_markup=kb)
            return
        if action == "pick" and value is not None:
            idx = int(value)
            if idx < 0 or idx >= len(items):
                await q.answer()
                return
            element = items[idx]
            st["element"] = element
            # готовим дефекты по связке этап+элемент
            stage = st.get("stage")
            defects = sorted(CNIL_CATALOG.get(stage, {}).get(element, {}).keys())
            st["items"]["defects"] = defects
            st["defects_selected"] = set()
            st["page"] = 0
            st["step"] = "defects"
            context.user_data["cnil"] = st

            kb = cnil_kb_multi(defects, st["defects_selected"], prefix="defects", page=0)
            await q.edit_message_text(
                f"Выбран элемент: {element}\n\n4/6. Выберите дефекты (можно несколько):",
                reply_markup=kb,
            )
            return

    # ---------- ДЕФЕКТЫ (мульти) ----------
    if scope == "defects":
        items = st.get("items", {}).get("defects", [])
        selected: set[int] = st.get("defects_selected") or set()
        if action == "page" and value is not None:
            page = int(value)
            kb = cnil_kb_multi(items, selected, prefix="defects", page=page)
            await q.edit_message_reply_markup(reply_markup=kb)
            return
        if action == "toggle" and value is not None:
            idx = int(value)
            if 0 <= idx < len(items):
                if idx in selected:
                    selected.remove(idx)
                else:
                    selected.add(idx)
                st["defects_selected"] = selected
                context.user_data["cnil"] = st
            kb = cnil_kb_multi(items, selected, prefix="defects", page=st.get("page", 0))
            await q.edit_message_reply_markup(reply_markup=kb)
            return
        if action == "reset":
            st["defects_selected"] = set()
            context.user_data["cnil"] = st
            kb = cnil_kb_multi(items, set(), prefix="defects", page=0)
            await q.edit_message_reply_markup(reply_markup=kb)
            return
        if action == "back":
            st["step"] = "element"
            stage = st.get("stage")
            elements = sorted(CNIL_CATALOG.get(stage, {}).keys())
            st["items"]["element"] = elements
            st["page"] = 0
            context.user_data["cnil"] = st

            kb = cnil_kb_list(elements, prefix="element", page=0)
            await q.edit_message_text(
                f"Выбран этап: {stage}\n\n3/6. Выберите конструктивный элемент:",
                reply_markup=kb,
            )
            return
        if action == "next":
            if not selected:
                await q.answer("Выберите хотя бы один дефект", show_alert=True)
                return
            # формируем список работ как объединение по выбранным дефектам
            stage = st.get("stage")
            element = st.get("element")
            defect_names = [items[i] for i in sorted(selected)]
            works_set = set()
            for dname in defect_names:
                works_set.update(CNIL_CATALOG.get(stage, {}).get(element, {}).get(dname, []))
            works = sorted(list(works_set))

            st["items"]["works"] = works
            st["works_selected"] = set()
            st["step"] = "works"
            st["page"] = 0
            context.user_data["cnil"] = st

            kb = cnil_kb_multi(works, set(), prefix="works", page=0)
            await q.edit_message_text("5/5. Выберите виды работ (можно несколько):", reply_markup=kb)
            return

    # ---------- РАБОТЫ (мульти) ----------
    if scope == "works":
        items = st.get("items", {}).get("works", [])
        selected: set[int] = st.get("works_selected") or set()
        if action == "page" and value is not None:
            page = int(value)
            kb = cnil_kb_multi(items, selected, prefix="works", page=page)
            await q.edit_message_reply_markup(reply_markup=kb)
            return
        if action == "toggle" and value is not None:
            idx = int(value)
            if 0 <= idx < len(items):
                if idx in selected:
                    selected.remove(idx)
                else:
                    selected.add(idx)
                st["works_selected"] = selected
                context.user_data["cnil"] = st
            kb = cnil_kb_multi(items, selected, prefix="works", page=st.get("page", 0))
            await q.edit_message_reply_markup(reply_markup=kb)
            return
        if action == "reset":
            st["works_selected"] = set()
            context.user_data["cnil"] = st
            kb = cnil_kb_multi(items, set(), prefix="works", page=0)
            await q.edit_message_reply_markup(reply_markup=kb)
            return
        if action == "back":
            st["step"] = "defects"
            context.user_data["cnil"] = st
            defects = st["items"].get("defects", [])
            kb = cnil_kb_multi(defects, st.get("defects_selected") or set(), prefix="defects", page=0)
            await q.edit_message_text("4/5. Выберите дефекты (можно несколько):", reply_markup=kb)
            return
        if action == "next":
            if not selected:
                await q.answer("Выберите хотя бы один вид работ", show_alert=True)
                return

            works = [items[i] for i in sorted(selected)]

            st["works_selected"] = selected
            st["work_marks"] = {}
            st["work_marks_order"] = works
            st["work_marks_idx"] = 0
            st["step"] = "work_marks"
            context.user_data["cnil"] = st

            first = works[0]
            await q.edit_message_text(
                f"Работа: {first}\nВведите «Отметка ОСИ, высотные отметки» (можно пропустить):",
                reply_markup=InlineKeyboardMarkup([[InlineKeyboardButton("✖️ Отмена", callback_data="cnil:cancel")]]),
            )
            return

    await q.answer()



async def error_handler(update: object, context: ContextTypes.DEFAULT_TYPE) -> None:
    # Prevent unhandled exceptions from crashing update processing
    logging.exception("Unhandled exception while handling update", exc_info=context.error)

def main():
    if not BOT_TOKEN:
        log.error("BOT_TOKEN не задан.")
        raise SystemExit("Укажите BOT_TOKEN в переменных окружения.")

    init_db()

    app = Application.builder().token(BOT_TOKEN).build()
    app.add_error_handler(error_handler)

    app.add_handler(CommandHandler("start", start))
    app.add_handler(CommandHandler("help", help_command))

    app.add_handler(CallbackQueryHandler(callback_handler))

    app.add_handler(MessageHandler(filters.VOICE, voice_router))

    app.add_handler(MessageHandler(filters.Document.ALL, document_handler))
    app.add_handler(MessageHandler(filters.TEXT & ~filters.COMMAND, text_router))

    log.info("Бот запущен...")
    app.run_polling()


if __name__ == "__main__":
    main()
